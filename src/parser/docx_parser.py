"""DOCX manuscript parsing for audiobook chapter extraction."""

from __future__ import annotations

import logging
import re
from collections import Counter
from dataclasses import dataclass, replace
from functools import lru_cache
from pathlib import Path
from typing import Any

from docx import Document as load_document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.opc.exceptions import PackageNotFoundError
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class BookMetadata:
    """Extracted title page metadata."""

    title: str
    subtitle: str | None
    author: str
    original_publisher: str | None


@dataclass(slots=True)
class Chapter:
    """Represents a single narratable chapter."""

    number: int
    title: str
    type: str
    raw_text: str
    word_count: int


@dataclass(slots=True)
class _ParagraphInfo:
    """Normalized paragraph data used during parsing."""

    index: int
    text: str
    style: str | None
    paragraph: Paragraph


class DocxParser:
    """Parse DOCX manuscripts into metadata and narratable chapters."""

    NON_AUTHOR_PHRASES: set[str] = {
        "a modern translation",
        "a new translation",
        "a contemporary translation",
        "a fresh translation",
        "modern translation",
        "new translation",
        "complete works",
        "selected works",
        "selected writings",
        "collected works",
        "collected writings",
        "essential writings",
        "classic edition",
        "modern edition",
        "revised edition",
        "new edition",
        "first edition",
        "second edition",
        "third edition",
        "annotated edition",
        "unabridged edition",
        "definitive edition",
        "library of alexandria",
        "table of contents",
        "introduction",
        "foreword",
        "preface",
        "prologue",
        "volume one",
        "volume two",
        "volume three",
        "volume i",
        "volume ii",
        "volume iii",
        "part one",
        "part two",
        "part three",
    }
    KNOWN_AUTHORS: dict[str, str] = {
        "sun tzu": "Sun Tzu",
        "marcus aurelius": "Marcus Aurelius",
        "lao tzu": "Lao Tzu",
        "miyamoto musashi": "Miyamoto Musashi",
        "james allen": "James Allen",
        "ralph waldo emerson": "Ralph Waldo Emerson",
        "emerson": "Ralph Waldo Emerson",
        "seneca": "Seneca",
        "senaca": "Seneca",
        "henry david thoreau": "Henry David Thoreau",
        "friedrich nietzsche": "Friedrich Nietzsche",
        "leo tolstoy": "Leo Tolstoy",
        "fyodor dostoevsky": "Fyodor Dostoevsky",
        "aristotle": "Aristotle",
        "epictetus": "Epictetus",
        "julius caesar": "Julius Caesar",
        "plato": "Plato",
        "homer": "Homer",
        "virgil": "Virgil",
        "ovid": "Ovid",
        "cicero": "Cicero",
        "thucydides": "Thucydides",
        "herodotus": "Herodotus",
        "confucius": "Confucius",
        "kierkegaard": "Søren Kierkegaard",
        "lucretius": "Lucretius",
        "paracelsus": "Paracelsus",
        "hermes trismegistus": "Hermes Trismegistus",
        "thoth hermes trismegistus": "Hermes Trismegistus",
        "napoleon": "Napoleon Bonaparte",
        "machiavelli": "Niccolò Machiavelli",
        "kant": "Immanuel Kant",
        "schopenhauer": "Arthur Schopenhauer",
        "voltaire": "Voltaire",
        "montaigne": "Michel de Montaigne",
        "thoreau": "Henry David Thoreau",
        "william walker atkinson": "William Walker Atkinson",
        "h g wells": "H.G. Wells",
        "jules verne": "Jules Verne",
        "edgar allan poe": "Edgar Allan Poe",
        "mary shelley": "Mary Shelley",
        "oscar wilde": "Oscar Wilde",
        "mark twain": "Mark Twain",
        "charles dickens": "Charles Dickens",
        "jane austen": "Jane Austen",
        "edgar rice burroughs": "Edgar Rice Burroughs",
        "arthur conan doyle": "Arthur Conan Doyle",
        "jack london": "Jack London",
        "robert louis stevenson": "Robert Louis Stevenson",
        "bram stoker": "Bram Stoker",
        "george macdonald": "George MacDonald",
        "samuel butler": "Samuel Butler",
        "charlotte perkins gilman": "Charlotte Perkins Gilman",
        "charlotte brontë": "Charlotte Brontë",
        "charlotte bronte": "Charlotte Brontë",
        "emily brontë": "Emily Brontë",
        "emily bronte": "Emily Brontë",
        "anne brontë": "Anne Brontë",
        "anne bronte": "Anne Brontë",
        "xenophon": "Xenophon",
        "sophocles": "Sophocles",
        "plutarch": "Plutarch",
        "apuleius": "Apuleius",
        "petronius": "Petronius",
        "lucian": "Lucian",
        "aristophanes": "Aristophanes",
        "cervantes": "Miguel de Cervantes",
        "miguel de cervantes": "Miguel de Cervantes",
        "alexis de tocqueville": "Alexis de Tocqueville",
        "tocqueville": "Alexis de Tocqueville",
        "karl marx": "Karl Marx",
        "carl von clausewitz": "Carl von Clausewitz",
        "clausewitz": "Carl von Clausewitz",
        "boethius": "Boethius",
        "guy de maupassant": "Guy de Maupassant",
        "maupassant": "Guy de Maupassant",
        "théophile gautier": "Théophile Gautier",
        "theophile gautier": "Théophile Gautier",
        "washington irving": "Washington Irving",
        "brothers grimm": "Brothers Grimm",
        "valmiki": "Valmiki",
        "vyasa": "Vyasa",
        "patanjali": "Patanjali",
        "zhuangzi": "Zhuangzi",
        "augustine": "Saint Augustine",
        "saint augustine": "Saint Augustine",
        "ignatius of loyola": "Ignatius of Loyola",
        "rené descartes": "René Descartes",
        "rene descartes": "René Descartes",
        "descartes": "René Descartes",
    }
    KNOWN_TITLES: dict[str, str] = {
        "arthashastra": "Kautilya",
        "instructions to his generals": "Frederick the Great",
        "history of the peloponnesian war": "Thucydides",
        "fear and trembling": "Søren Kierkegaard",
        "on the nature of things": "Lucretius",
        "the chaldean oracles": "Anonymous",
        "the book concerning the tincture of the philosophers": "Paracelsus",
        "on sense and the sensible": "Aristotle",
        "on life and death": "Aristotle",
        "on memory and reminiscence": "Aristotle",
        "on sleep and sleeplessness": "Aristotle",
        "on dreams": "Aristotle",
        "metaphysics": "Aristotle",
        "on longevity and shortness of life": "Aristotle",
        "rhetoric": "Aristotle",
        "the emerald tablet of thoth": "Hermes Trismegistus",
        "in the year 2889": "Jules Verne",
        "corpus hermeticum": "Hermes Trismegistus",
        "the emerald tablet": "Hermes Trismegistus",
        "the hermetic and alchemical writings of paracelsus": "Paracelsus",
        "coelum philosophorum": "Paracelsus",
        "the master key system": "Charles F. Haanel",
        "aristotles complete works on the mind dreams and the nature of thought": "Aristotle",
        "aristotle's metaphysical and scientific masterpieces": "Aristotle",
        "aristotle's insights into memory, sleep, and the mysteries of the human mind": "Aristotle",
        "the virgin of the world": "Hermes Trismegistus",
        "the life and teachings of thoth hermes trismegistus": "Hermes Trismegistus",
        "micromegas: a philosophical story": "Voltaire",
        "micromegas": "Voltaire",
        "candide": "Voltaire",
        "zadig": "Voltaire",
        "the wreck of the golden mary": "Charles Dickens",
        "the perils of certain english prisoners": "Charles Dickens",
        "the haunted house": "Charles Dickens",
        "the lazy tour of two idle apprentices": "Charles Dickens",
        "a house to let": "Charles Dickens",
        "no thoroughfare": "Charles Dickens",
        "the battle of life": "Charles Dickens",
        "jane eyre": "Charlotte Brontë",
        "villette": "Charlotte Brontë",
        "shirley": "Charlotte Brontë",
        "the professor": "Charlotte Brontë",
        "wuthering heights": "Emily Brontë",
        "the tenant of wildfell hall": "Anne Brontë",
        "agnes grey": "Anne Brontë",
        "the iliad: the fall of troy": "Homer",
        "the iliad": "Homer",
        "the odyssey": "Homer",
        "anabasis": "Xenophon",
        "memorabilia": "Xenophon",
        "the last days of socrates": "Plato",
        "euthyphro": "Plato",
        "apology": "Plato",
        "crito": "Plato",
        "phaedo": "Plato",
        "the three theban plays": "Sophocles",
        "oedipus the king": "Sophocles",
        "oedipus at colonus": "Sophocles",
        "antigone": "Sophocles",
        "meno": "Plato",
        "symposium": "Plato",
        "the republic": "Plato",
        "true history": "Lucian",
        "parallel lives": "Plutarch",
        "the golden ass": "Apuleius",
        "the satyricon": "Petronius",
        "clouds": "Aristophanes",
        "oeconomicus": "Xenophon",
        "apology of socrates": "Xenophon",
        "plato apology": "Plato",
        "charmides": "Plato",
        "plato lesser hippias": "Plato",
        "plato minos": "Plato",
        "plato clitophon": "Plato",
        "plato epinomis": "Plato",
        "plato cratylus": "Plato",
        "plato alcibiades ii": "Plato",
        "plato theages": "Plato",
        "socrates axiochus": "Plato",
        "axiochus": "Plato",
        "socrates eryxias": "Plato",
        "eryxias": "Plato",
        "socrates demodocus": "Plato",
        "demodocus": "Plato",
        "the bhagavad gita": "Vyasa",
        "the dhammapada": "Buddha",
        "the tao of chuang tzu (zhuangzi)": "Zhuangzi",
        "the tao of chuang tzu": "Zhuangzi",
        "the yoga sutras of patanjali": "Patanjali",
        "the diamond sutra": "Anonymous",
        "the ramayana": "Valmiki",
        "confessions": "Saint Augustine",
        "the spiritual exercises": "Ignatius of Loyola",
        "the book of tea": "Kakuzo Okakura",
        "the pillow book": "Sei Shōnagon",
        "meditations of descartes": "René Descartes",
        "el cantar de mio cid": "Anonymous",
        "the song of roland": "Anonymous",
        "the egyptian book of the dead": "Anonymous",
        "the tibetan book of the dead (bardo thodol)": "Padmasambhava",
        "the tibetan book of the dead": "Padmasambhava",
        "the popol vuh": "Anonymous",
        "the epic of gilgamesh": "Anonymous",
        "the descent of ishtar": "Anonymous",
        "nergal and ereshkigal": "Anonymous",
        "the lament for ur": "Anonymous",
        "the marriage of martu": "Anonymous",
        "temple hymn to nanna (sin)": "Enheduanna",
        "temple hymn to nanna": "Enheduanna",
        "the debate between sheep and grain": "Anonymous",
        "erra and ishum": "Kabti-ilani-Marduk",
        "the first book of enoch": "Anonymous",
        "the first book of maccabees": "Anonymous",
        "the second book of esdras": "Anonymous",
        "the second book of the maccabees": "Anonymous",
        "the third book of maccabees": "Anonymous",
        "the fourth book of maccabees": "Anonymous",
        "the book of jubilees": "Anonymous",
        "the testament of the twelve patriarchs": "Anonymous",
        "the book of baruch": "Baruch ben Neriah",
        "the third book of baruch": "Anonymous",
        "the apocalypse of peter": "Anonymous",
        "the book of the secrets of enoch": "Anonymous",
        "the hebrew book of enoch": "Anonymous",
        "the acts of paul and thecla": "Anonymous",
        "the book of the watchers": "Anonymous",
        "pistis sophia": "Anonymous",
        "the gospel of philip": "Anonymous",
        "the gospel of thomas": "Anonymous",
        "the gospel of judas": "Anonymous",
        "the book of thomas the contender": "Anonymous",
        "the apocryphon of john (the secret book of john)": "Anonymous",
        "the apocryphon of john": "Anonymous",
        "the apocalypse of abraham": "Anonymous",
        "the sophia of jesus christ": "Anonymous",
        "the words of gad the seer": "Anonymous",
        "the gospel of the egyptians": "Anonymous",
        "thunder, perfect mind": "Anonymous",
        "the dialogue of the savior": "Anonymous",
        "the odes of solomon": "Anonymous",
        "the psalms of solomon": "Anonymous",
        "the epistle of barnabas": "Anonymous",
        "the book of creation (sefer yetzirah)": "Anonymous",
        "the book of creation": "Anonymous",
        "the gospel of the hebrews": "Anonymous",
        "the gospel of the nazarenes": "Anonymous",
        "the epistle to the laodiceans": "Anonymous",
        "the book of enoch & the fallen angels": "Anonymous",
        "the book of enoch and the fallen angels": "Anonymous",
        "the books of baruch & the exiles": "Anonymous",
        "the books of baruch and the exiles": "Anonymous",
        "the nag hammadi scriptures": "Anonymous",
        "the shepherd of hermas": "Hermas",
        "ecclesiasticus (the wisdom of jesus the son of sirach)": "Ben Sira",
        "ecclesiasticus": "Ben Sira",
        "the second treatise of the great seth": "Anonymous",
        "the acts of peter and the twelve apostles": "Anonymous",
        "the dead sea scrolls bible": "Anonymous",
        "don quixote": "Miguel de Cervantes",
        "the communist manifesto": "Karl Marx",
        "on war": "Carl von Clausewitz",
        "the federalist papers": "Alexander Hamilton, James Madison & John Jay",
        "the us constitution": "United States Founding Fathers",
        "the declaration of independence": "Thomas Jefferson",
        "the bill of rights n constitutional amendments": "United States Congress",
        "the bill of rights and constitutional amendments": "United States Congress",
        "the constitution of the us n declaration of independence": "United States Founding Fathers",
        "the constitution of the us and declaration of independence": "United States Founding Fathers",
        "constitution of the united states pocket edition": "United States Founding Fathers",
        "constitution of the united states large print": "United States Founding Fathers",
        "the simple sabotage field manual": "Office of Strategic Services",
        "democracy in america": "Alexis de Tocqueville",
        "the complete essays": "Michel de Montaigne",
        "theological tractates": "Boethius",
        "on the trinity (de trinitate)": "Boethius",
        "on the trinity": "Boethius",
        "grimms' fairy tales (complete)": "Brothers Grimm",
        "grimms' fairy tales": "Brothers Grimm",
        "unveiling a parallel": "Alice Ilgenfritz Jones",
        "varney the vampire": "James Malcolm Rymer",
        "the horla": "Guy de Maupassant",
        "the mummy's foot": "Théophile Gautier",
        "the mummys foot": "Théophile Gautier",
        "rip van winkle": "Washington Irving",
        "the unparalleled adventure of one hans pfaall": "Edgar Allan Poe",
        "the complete ethiopian bible (volume i)": "Anonymous",
        "the complete ethiopian bible (volume ii)": "Anonymous",
        "the complete ethiopian bible": "Anonymous",
        "the ultimate horror collection of 101+ macabre masterpieces (volume i)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume ii)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume iii)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume iv)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume vi)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume vii)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume viii)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume ix)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume x)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume xi)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume xii)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume xiii)": "Various Authors",
        "the ultimate horror collection of 101+ macabre masterpieces (volume xiv)": "Various Authors",
        "the complete strategy collection (volume i)": "Various Authors",
        "the complete strategy collection (volume ii)": "Various Authors",
        "the complete strategy collection (volume iii)": "Various Authors",
        "the complete philosophy collection (vol. i)": "Various Authors",
        "the complete philosophy collection (vol. iii)": "Various Authors",
        "the complete philosophy collection (vol. iv)": "Various Authors",
        "the complete philosophy collection (vol. v)": "Various Authors",
        "the complete philosophy collection (vol. vi)": "Various Authors",
        "the complete philosophy collection (vol. viii)": "Various Authors",
        "the complete leadership collection": "Various Authors",
        "the complete leadership collection (volume i)": "Various Authors",
        "the complete leadership collection (volume ii)": "Various Authors",
        "the complete hermeticism philosophy collection": "Various Authors",
        "the ancient wisdom collection": "Various Authors",
        "the ultimate esoteric wisdom collection": "Various Authors",
        "the classical wisdom collection": "Various Authors",
        "the complete strategy & war collection": "Various Authors",
        "the complete strategy and war collection": "Various Authors",
        "the complete warrior's mindset collection": "Various Authors",
        "the complete warriors mindset collection": "Various Authors",
        "mystical hermetic & christian dialogues": "Various Authors",
        "mystical hermetic and christian dialogues": "Various Authors",
    }

    def __init__(self) -> None:
        """Initialize chapter detection and skip rules."""

        self.chapter_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(
                r"^chapter\s+(?P<number>[ivxlcdm]+|\d+|[a-z][a-z-]*)\s*[:.\-\u2013\u2014]?\s*(?P<title>.*)$",
                re.IGNORECASE,
            ),
            re.compile(r"^(?P<number>[ivxlcdm]+|\d+)\s*[:.\-]\s*(?P<title>.+)$", re.IGNORECASE),
        )
        self.book_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(
                r"^book\s+(?:the\s+)?(?P<number>[ivxlcdm]+|\d+|[a-z][a-z-]*)(?:\s*[:.\-]?\s*(?P<title>.*))?$",
                re.IGNORECASE,
            ),
            re.compile(
                r"^(?P<number>first|second|third|fourth|fifth|sixth|seventh|eighth|ninth|tenth"
                r"|eleventh|twelfth|thirteenth|fourteenth|fifteenth|sixteenth|seventeenth"
                r"|eighteenth|nineteenth|twentieth)\s+book\b(?:\s*[:.\-–—\s]\s*(?P<title>.*))?$",
                re.IGNORECASE,
            ),
        )
        self.part_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(
                r"^part\s+(?:the\s+)?(?P<number>[ivxlcdm]+|\d+|[a-z][a-z-]*)(?:\s*[:.\-–—]?\s*(?P<title>.*))?$",
                re.IGNORECASE,
            ),
        )
        self.maxim_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(
                r"^maxim\s+(?P<number>[ivxlcdm]+|\d+)\b(?:\s*[:.\-]?\s*(?P<title>.*))?$",
                re.IGNORECASE,
            ),
        )
        self.intro_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(
                r"^(?P<label>introduction|preface|prologue)\b(?:\s*[:.\-]\s*(?P<title>.*))?$",
                re.IGNORECASE,
            ),
        )
        reader_note_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(r"^preface(?:\s*[—-]\s*|\s+)message to the reader\b", re.IGNORECASE),
            re.compile(r"^message\s+to\s+the\s+reader\b", re.IGNORECASE),
            re.compile(r"^(?:a\s+)?note\s+to\s+the\s+reader\b", re.IGNORECASE),
        )
        outro_note_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(r"^outro\b", re.IGNORECASE),
            re.compile(r"^closing\s+note\b", re.IGNORECASE),
            re.compile(r"^afterword\b", re.IGNORECASE),
            re.compile(r"^(?:a\s+word\s+from\s+the\s+author|from\s+the\s+author)\b", re.IGNORECASE),
            re.compile(r"^author(?:s|\s+s)?\s+note\b", re.IGNORECASE),
        )
        self.front_matter_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(r"^(copyright|©)\b", re.IGNORECASE),
            re.compile(r"^(table of contents|contents)\b", re.IGNORECASE),
            re.compile(r"^(about\s+the\s+author|about\s+this\s+translation|translator.?s\s+note)\b", re.IGNORECASE),
            re.compile(r"^foreword\b", re.IGNORECASE),
            *reader_note_patterns,
            *outro_note_patterns,
        )
        self.back_matter_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(r"^thank(?:s| you)?\s+you\s+for\s+reading\b", re.IGNORECASE),
            re.compile(r"^thank\s+for\s+reading\b", re.IGNORECASE),
            re.compile(r"^epilogue\b", re.IGNORECASE),
            re.compile(r"^the\s+end(?:\s*[.!?])?\s*$", re.IGNORECASE),
            *reader_note_patterns,
            *outro_note_patterns,
        )
        self.author_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(r"^(?:by|written by)\s+(?P<author>.+)$", re.IGNORECASE),
        )
        self.publisher_patterns: tuple[re.Pattern[str], ...] = (
            re.compile(r"^(?:originally\s+published\s+by|published\s+by)\s+(?P<publisher>.+)$", re.IGNORECASE),
        )
        self.word_number_map: dict[str, int] = {
            "one": 1,
            "two": 2,
            "three": 3,
            "four": 4,
            "five": 5,
            "six": 6,
            "seven": 7,
            "eight": 8,
            "nine": 9,
            "ten": 10,
            "eleven": 11,
            "twelve": 12,
            "thirteen": 13,
            "fourteen": 14,
            "fifteen": 15,
            "sixteen": 16,
            "seventeen": 17,
            "eighteen": 18,
            "nineteen": 19,
            "twenty": 20,
            "first": 1,
            "second": 2,
            "third": 3,
            "fourth": 4,
            "fifth": 5,
            "sixth": 6,
            "seventh": 7,
            "eighth": 8,
            "ninth": 9,
            "tenth": 10,
            "eleventh": 11,
            "twelfth": 12,
            "thirteenth": 13,
            "fourteenth": 14,
            "fifteenth": 15,
            "sixteenth": 16,
            "seventeenth": 17,
            "eighteenth": 18,
            "nineteenth": 19,
            "twentieth": 20,
        }
        self.last_toc_entries: list[str] = []

    def parse(self, docx_path: str | Path) -> tuple[BookMetadata, list[Chapter]]:
        """
        Parse a DOCX file and return extracted metadata plus chapter content.

        Args:
            docx_path: Path to the DOCX manuscript.

        Returns:
            A tuple of book metadata and detected chapters.

        Raises:
            ValueError: If the file is missing, unreadable, or missing critical metadata.
        """

        path = Path(docx_path)
        if not path.exists():
            raise ValueError(f"DOCX file does not exist: {path}")

        logger.info("Parsing DOCX manuscript: %s", path)

        try:
            document = load_document(str(path))
        except PackageNotFoundError as exc:
            raise ValueError(f"Invalid DOCX file: {path}") from exc
        except Exception as exc:
            raise ValueError(f"Failed to read DOCX file: {path}") from exc

        metadata = self._extract_metadata(document)
        chapters = self._extract_chapters(document, metadata.title, metadata.author)
        if not chapters:
            chapters = self._fallback_single_chapter(document, path, metadata_title=metadata.title)
        if not chapters:
            total_paragraphs = len(document.paragraphs)
            non_empty_paragraphs = sum(1 for paragraph in document.paragraphs if paragraph.text.strip())
            raise ValueError(
                f"No narratable chapters detected in {path.name}. "
                f"The document has {total_paragraphs} paragraphs ({non_empty_paragraphs} non-empty). "
                "The parser requires chapter headings with 'Chapter N' format or Heading-styled paragraphs."
            )
        if len(chapters) == 1 and chapters[0].title == "Full Text":
            logger.warning(
                "No chapter headings found in %s; using whole-document fallback (%s words)",
                path.name,
                chapters[0].word_count,
            )

        logger.info("Extracted %s narratable chapters from %s", len(chapters), path)
        return metadata, chapters

    def parse_with_folder_hint(
        self,
        docx_path: str | Path,
        folder_name: str | None = None,
    ) -> tuple[BookMetadata, list[Chapter]]:
        """Parse a DOCX file, using the folder name as an author hint if needed."""

        metadata, chapters = self.parse(docx_path)

        if metadata.author == "Unknown Author":
            if folder_name:
                folder_author = self._extract_author_from_folder(folder_name)
                if folder_author:
                    logger.info("Using folder-name author hint: %s", folder_author)
                    metadata = replace(metadata, author=folder_author)

            if metadata.author == "Unknown Author" and metadata.title:
                known_author = None
                for title_lower in self._title_lookup_candidates(metadata.title):
                    known_author = self.KNOWN_TITLES.get(title_lower)
                    if known_author and known_author != "Unknown Author":
                        break
                if known_author and known_author != "Unknown Author":
                    logger.info("Using title-based author lookup: %s", known_author)
                    metadata = replace(metadata, author=known_author)

        return metadata, chapters

    def _title_lookup_candidates(self, title: str) -> list[str]:
        """Return normalized title variants for KNOWN_TITLES lookup."""

        base = self._normalize_text(title).casefold()
        variants = [
            base,
            base.replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"'),
        ]
        transforms = (
            lambda value: re.sub(r"\s*&\s*", " and ", value),
            lambda value: re.sub(r"\bn\b", "and", value),
            lambda value: re.sub(r"\s*\([^)]*\)", "", value),
        )

        ordered: list[str] = []
        seen: set[str] = set()
        queue = list(variants)

        while queue:
            candidate = self._normalize_text(queue.pop(0))
            if not candidate or candidate in seen:
                continue
            seen.add(candidate)
            ordered.append(candidate)
            for transform in transforms:
                transformed = self._normalize_text(transform(candidate))
                if transformed and transformed not in seen:
                    queue.append(transformed)

        return ordered

    def _fallback_single_chapter(
        self,
        doc: DocxDocument,
        path: Path,
        *,
        metadata_title: str | None = None,
    ) -> list[Chapter]:
        """Return a single fallback chapter when no headings are detected."""

        del path
        paragraphs = self._collect_paragraphs(doc)
        running_headers = self._detect_running_headers(paragraphs, metadata_title=metadata_title)
        body_started = False
        collecting_toc = False
        body_paragraphs: list[str] = []

        for paragraph in paragraphs:
            text = paragraph.text
            style = paragraph.style

            if self._is_toc_heading(text):
                collecting_toc = True
                continue

            if collecting_toc:
                if self._looks_like_toc_entry(text, style) or (style and "toc" in style.lower()):
                    continue
                collecting_toc = False

            if self._is_back_matter_section(text):
                if body_started:
                    break
                continue

            if self._is_skip_section(text) or self._looks_like_credit_or_note(text):
                continue

            if not body_started:
                if self._count_words(text) <= 30:
                    continue
                body_started = True

            if self._is_running_header_paragraph(
                text,
                style,
                running_headers,
                metadata_title=metadata_title,
            ):
                logger.debug("Skipping running header '%s' in fallback chapter", text)
                continue

            body_paragraphs.append(text)

        fallback = self._build_chapter(
            {"number": 1, "title": "Full Text", "type": "chapter"},
            body_paragraphs,
        )
        return [fallback] if fallback is not None else []

    def _extract_metadata(self, doc: DocxDocument) -> BookMetadata:
        """Extract title-page metadata from the first section of the document."""

        paragraphs = self._collect_paragraphs(doc, limit=20)
        if not paragraphs:
            raise ValueError("Document does not contain readable text for metadata extraction.")

        title_position, title = self._find_title(paragraphs)
        author_position, author = self._find_author(paragraphs, title_position)
        subtitle = self._find_subtitle(paragraphs, title_position, author_position)
        original_publisher = self._find_original_publisher(paragraphs)

        return BookMetadata(
            title=title,
            subtitle=subtitle,
            author=author,
            original_publisher=original_publisher,
        )

    def _extract_chapters(
        self,
        doc: DocxDocument,
        metadata_title: str | None = None,
        metadata_author: str | None = None,
    ) -> list[Chapter]:
        """Extract narratable introduction and chapter bodies from the document."""

        paragraphs = self._collect_paragraphs(doc)
        chapters: list[Chapter] = []
        current_heading: dict[str, Any] | None = None
        current_heading_style: str | None = None
        current_body: list[str] = []
        narration_started = False
        real_chapter_started = False
        saw_numbered_chapter = False
        collecting_toc = False
        skipping_prefatory_section = False
        next_chapter_number = 1
        primary_numbered_heading_level = self._detect_primary_numbered_heading_level(paragraphs)
        has_standalone_chapter_markers = any(
            self._standalone_chapter_marker_number(paragraph.text, paragraph.style) is not None
            for paragraph in paragraphs
        )
        used_dominant_heading_fallback = False
        if primary_numbered_heading_level is None and not has_standalone_chapter_markers:
            primary_numbered_heading_level = self._detect_dominant_heading_level(paragraphs)
            used_dominant_heading_fallback = primary_numbered_heading_level is not None
        author_key = self._normalize_text(metadata_author).casefold() if metadata_author else None
        running_headers = self._detect_running_headers(
            paragraphs,
            metadata_title=metadata_title,
        )
        pending_standalone_chapter_number: int | None = None
        self.last_toc_entries = []

        for paragraph in paragraphs:
            text = paragraph.text
            style = paragraph.style

            if not narration_started and self._is_toc_heading(text):
                collecting_toc = True
                logger.debug("Detected TOC heading at paragraph %s", paragraph.index)
                continue

            if collecting_toc:
                if self._looks_like_chapter_style(style):
                    collecting_toc = False
                elif self._looks_like_toc_entry(text, style):
                    self.last_toc_entries.append(text)
                    continue
                elif style and "toc" in style.lower():
                    self.last_toc_entries.append(text)
                    continue
                else:
                    collecting_toc = False
                    continue

            if current_heading is not None and self._is_running_header_paragraph(
                text,
                style,
                running_headers,
                metadata_title=metadata_title,
                current_heading_title=current_heading["title"],
            ):
                if (
                    not current_body
                    and author_key
                    and text.casefold() == author_key
                    and current_heading_style is not None
                    and self._looks_like_primary_section_style(current_heading_style)
                ):
                    pass
                else:
                    logger.debug("Skipping running header '%s' in chapter '%s'", text, current_heading["title"])
                    continue

            standalone_chapter_number = self._standalone_chapter_marker_number(text, style)
            if standalone_chapter_number is not None:
                if current_heading is not None:
                    built = self._build_chapter(current_heading, current_body)
                    if built is not None:
                        chapters.append(built)
                    current_heading = None
                    current_heading_style = None
                    current_body = []
                pending_standalone_chapter_number = standalone_chapter_number
                narration_started = True
                real_chapter_started = True
                saw_numbered_chapter = True
                next_chapter_number = max(next_chapter_number, standalone_chapter_number + 1)
                logger.debug(
                    "Captured standalone chapter marker at paragraph %s: Chapter %s",
                    paragraph.index,
                    standalone_chapter_number,
                )
                continue

            explicit_heading, parsed_heading = self._is_chapter_heading(
                text,
                style,
                primary_numbered_heading_level=primary_numbered_heading_level,
            )
            is_heading = explicit_heading
            if not is_heading:
                is_heading, parsed_heading = self._is_implicit_section_heading(
                    text,
                    style,
                    chapter_number=next_chapter_number,
                    saw_numbered_chapter=saw_numbered_chapter,
                    primary_numbered_heading_level=primary_numbered_heading_level,
                    allow_same_level=used_dominant_heading_fallback,
                    pending_standalone_chapter_number=pending_standalone_chapter_number,
                )
            is_skip_section = self._is_skip_section(text)
            is_back_matter_section = self._is_back_matter_section(text)

            if skipping_prefatory_section:
                if standalone_chapter_number is not None:
                    skipping_prefatory_section = False
                elif is_heading and parsed_heading is not None and not is_skip_section and not is_back_matter_section:
                    skipping_prefatory_section = False
                else:
                    logger.debug(
                        "Skipping non-narrated prefatory paragraph %s: %s",
                        paragraph.index,
                        text,
                    )
                    continue

            if not real_chapter_started and is_skip_section:
                if current_heading is not None:
                    built = self._build_chapter(current_heading, current_body)
                    if built is not None:
                        chapters.append(built)
                    current_heading = None
                    current_heading_style = None
                    current_body = []
                pending_standalone_chapter_number = None
                skipping_prefatory_section = True
                logger.debug("Skipping pre-chapter note section at paragraph %s: %s", paragraph.index, text)
                continue

            if not narration_started:
                if is_skip_section:
                    logger.debug("Skipping front matter paragraph %s: %s", paragraph.index, text)
                    continue
                if pending_standalone_chapter_number is not None:
                    narration_started = True
                    current_heading = {
                        "number": pending_standalone_chapter_number,
                        "title": f"Chapter {pending_standalone_chapter_number}",
                        "type": "chapter",
                    }
                    current_heading_style = style
                    current_body = []
                    real_chapter_started = True
                    pending_standalone_chapter_number = None
                    logger.debug(
                        "Started chapter from standalone marker before paragraph %s: %s",
                        paragraph.index,
                        current_heading["title"],
                    )
                if not narration_started:
                    if not is_heading or parsed_heading is None:
                        continue

                    narration_started = True
                    current_heading = parsed_heading
                    current_heading_style = style
                    current_body = []
                    if parsed_heading["type"] == "chapter":
                        real_chapter_started = True
                        if explicit_heading:
                            saw_numbered_chapter = True
                        next_chapter_number = max(next_chapter_number, parsed_heading["number"] + 1)
                    pending_standalone_chapter_number = None
                    logger.debug(
                        "Started narratable section at paragraph %s: %s",
                        paragraph.index,
                        parsed_heading["title"],
                    )
                    continue

            if is_heading and parsed_heading is not None:
                pending_standalone_chapter_number = None
                if parsed_heading["type"] == "introduction":
                    if current_heading is None:
                        current_heading = parsed_heading
                        current_heading_style = style
                        current_body = []
                    else:
                        current_body.append(text)
                    continue

                if (
                    current_heading is not None
                    and current_heading["type"] == "introduction"
                    and parsed_heading["type"] != "chapter"
                    and not self._is_explicit_chapter_heading(text)
                    and not self._looks_like_primary_section_style(style)
                ):
                    current_body.append(text)
                    continue

                if current_heading is None:
                    current_heading = parsed_heading
                    current_heading_style = style
                    current_body = []
                else:
                    built = self._build_chapter(current_heading, current_body)
                    if built is not None:
                        chapters.append(built)
                    current_heading = parsed_heading
                    current_heading_style = style
                    current_body = []
                if parsed_heading["type"] == "chapter":
                    real_chapter_started = True
                    if explicit_heading:
                        saw_numbered_chapter = True
                    next_chapter_number = max(next_chapter_number, parsed_heading["number"] + 1)
                logger.debug(
                    "Detected chapter %s at paragraph %s",
                    parsed_heading["number"],
                    paragraph.index,
                )
                continue

            if is_back_matter_section:
                logger.debug("Reached back matter at paragraph %s", paragraph.index)
                break

            if is_skip_section:
                logger.debug("Skipping non-narrated paragraph %s after start: %s", paragraph.index, text)
                continue

            if self._looks_like_chapter_style(style) and self._looks_like_generic_section_divider(text):
                logger.debug("Skipping generic section divider at paragraph %s: %s", paragraph.index, text)
                continue

            if pending_standalone_chapter_number is not None and current_heading is None:
                if (
                    self._looks_like_chapter_style(style)
                    and not self._looks_like_generic_section_divider(text)
                    and not self._looks_like_author_heading_metadata(text)
                    and not self._looks_like_credit_or_note(text)
                ):
                    current_heading = {
                        "number": pending_standalone_chapter_number,
                        "title": text,
                        "type": "chapter",
                    }
                    current_heading_style = style
                    current_body = []
                    real_chapter_started = True
                    pending_standalone_chapter_number = None
                    logger.debug(
                        "Promoted standalone marker to titled heading before paragraph %s: %s",
                        paragraph.index,
                        current_heading["title"],
                    )
                    continue
                current_heading = {
                    "number": pending_standalone_chapter_number,
                    "title": f"Chapter {pending_standalone_chapter_number}",
                    "type": "chapter",
                }
                current_heading_style = style
                current_body = []
                real_chapter_started = True
                pending_standalone_chapter_number = None
                logger.debug(
                    "Started body-driven chapter from standalone marker before paragraph %s: %s",
                    paragraph.index,
                    current_heading["title"],
                )

            if current_heading is not None:
                if current_heading["type"] == "chapter" and not current_body:
                    if (
                        self._is_generic_chapter_title(current_heading["title"])
                        and self._looks_like_followup_chapter_title(text, allow_long=True)
                    ):
                        logger.debug(
                            "Replacing generic chapter title '%s' with follow-up title at paragraph %s: %s",
                            current_heading["title"],
                            paragraph.index,
                            text,
                        )
                        current_heading["title"] = text
                        continue

                    if (
                        current_heading_style is not None
                        and self._looks_like_chapter_style(current_heading_style)
                        and not self._is_generic_chapter_title(current_heading["title"])
                        and len(text.split()) >= 2
                        and self._looks_like_followup_chapter_title(text, allow_long=False)
                    ):
                        logger.debug(
                            "Appending follow-up title fragment to '%s' at paragraph %s: %s",
                            current_heading["title"],
                            paragraph.index,
                            text,
                        )
                        current_heading["title"] = f"{current_heading['title']} {text}"
                        continue

                current_body.append(text)

        if current_heading is not None:
            built = self._build_chapter(current_heading, current_body)
            if built is not None:
                chapters.append(built)

        if saw_numbered_chapter:
            self._validate_toc(chapters)

        return chapters

    def _detect_running_headers(
        self,
        paragraphs: list[_ParagraphInfo],
        *,
        metadata_title: str | None = None,
    ) -> set[str]:
        """Detect repeated Normal-style phrases that behave like running headers."""

        repeated_normal_phrases: Counter[str] = Counter()
        normalized_title = self._normalize_text(metadata_title).casefold() if metadata_title else None

        for paragraph in paragraphs:
            if not paragraph.style or paragraph.style.casefold() != "normal":
                continue
            normalized_text = paragraph.text.casefold()
            word_count = len(paragraph.text.split())
            if word_count == 0:
                continue
            if word_count <= 5:
                repeated_normal_phrases[normalized_text] += 1
                continue
            if normalized_title and (
                normalized_text == normalized_title
                or normalized_text.startswith(f"{normalized_title}:")
                or normalized_text.startswith(f"{normalized_title} ")
            ):
                repeated_normal_phrases[normalized_text] += 1
                continue
            if ":" in paragraph.text and word_count <= 10:
                repeated_normal_phrases[normalized_text] += 1

        return {phrase for phrase, count in repeated_normal_phrases.items() if count >= 3}

    def _is_implicit_section_heading(
        self,
        text: str,
        style: str | None,
        chapter_number: int,
        *,
        saw_numbered_chapter: bool,
        primary_numbered_heading_level: int | None,
        allow_same_level: bool = False,
        pending_standalone_chapter_number: int | None = None,
    ) -> tuple[bool, dict[str, Any] | None]:
        """Treat level-one heading titles as chapter boundaries when numbering is absent."""

        normalized_text = self._normalize_text(text)
        if not normalized_text:
            return False, None
        heading_level = self._heading_level(style)
        if heading_level is None:
            return False, None
        if primary_numbered_heading_level is None:
            if not self._looks_like_primary_section_style(style):
                return False, None
        else:
            if heading_level > primary_numbered_heading_level:
                return False, None
            if heading_level == 1 and primary_numbered_heading_level > 1:
                return False, None
            if allow_same_level and heading_level != primary_numbered_heading_level:
                return False, None
            if (
                heading_level == primary_numbered_heading_level
                and not allow_same_level
                and pending_standalone_chapter_number is None
            ):
                return False, None
            if (
                heading_level < primary_numbered_heading_level
                and pending_standalone_chapter_number is None
            ):
                if saw_numbered_chapter and (
                    not self._looks_like_primary_section_style(style)
                    or self._looks_like_generic_section_divider(normalized_text)
                ):
                    return False, None
        if self._is_skip_section(normalized_text) or self._is_back_matter_section(normalized_text):
            return False, None
        if self._looks_like_author_heading_metadata(normalized_text):
            return False, None
        if self._looks_like_credit_or_note(normalized_text):
            return False, None
        if len(normalized_text.split()) > 18:
            return False, None
        if re.search(r"[,;:]$", normalized_text):
            return False, None
        return True, {
            "number": pending_standalone_chapter_number or chapter_number,
            "title": normalized_text,
            "type": "chapter",
        }

    def _is_explicit_chapter_heading(self, text: str) -> bool:
        """Return whether text is an explicit 'Chapter N' heading."""

        normalized_text = self._normalize_text(text)
        return bool(normalized_text and self.chapter_patterns[0].match(normalized_text))

    def _is_chapter_heading(
        self,
        text: str,
        style: str | None,
        *,
        primary_numbered_heading_level: int | None = None,
    ) -> tuple[bool, dict[str, Any] | None]:
        """
        Determine whether a paragraph is a chapter or introduction heading.

        Returns:
            A `(bool, parsed_data)` tuple, where parsed data contains `number`, `title`, and `type`.
        """

        normalized_text = self._normalize_text(text)
        if not normalized_text:
            return False, None
        if self._is_skip_section(normalized_text):
            return False, None
        if self._looks_like_author_heading_metadata(normalized_text):
            return False, None

        style_is_heading = self._looks_like_chapter_style(style)
        heading_level = self._heading_level(style)

        if (
            primary_numbered_heading_level is not None
            and heading_level is not None
            and heading_level > primary_numbered_heading_level
        ):
            return False, None

        for pattern in self.intro_patterns:
            match = pattern.match(normalized_text)
            intro_label = self._normalize_text(match.group("label")).casefold() if match else ""
            if match and (style_is_heading or intro_label in {"introduction", "prologue"}):
                title = self._normalize_text(match.group("label"))
                return True, {"number": 0, "title": title.title(), "type": "introduction"}

        for pattern in self.chapter_patterns:
            match = pattern.match(normalized_text)
            if not match:
                continue

            if primary_numbered_heading_level is not None and not style_is_heading:
                continue
            if pattern is self.chapter_patterns[1] and not style_is_heading:
                continue

            chapter_number = self._coerce_chapter_number(match.group("number"))
            if chapter_number is None:
                continue

            title = self._normalize_text(match.group("title"))
            if pattern is self.chapter_patterns[1] and not title:
                continue

            return True, {
                "number": chapter_number,
                "title": title or f"Chapter {chapter_number}",
                "type": "chapter",
            }

        for pattern in self.book_patterns + self.part_patterns + self.maxim_patterns:
            match = pattern.match(normalized_text)
            if not match:
                continue
            if not style_is_heading:
                continue

            chapter_number = self._coerce_chapter_number(match.group("number"))
            if chapter_number is None:
                continue

            return True, {
                "number": chapter_number,
                "title": normalized_text,
                "type": "chapter",
            }

        return False, None

    def _is_skip_section(self, text: str) -> bool:
        """Return whether a section should be skipped for narration."""

        normalized_text = self._normalize_heading_for_skip_rules(text)
        return any(pattern.match(normalized_text) for pattern in self.front_matter_patterns + self.back_matter_patterns)

    def _count_words(self, text: str) -> int:
        """Count words in the provided text."""

        return len(re.findall(r"\b[\w']+\b", text))

    def _standalone_chapter_marker_number(self, text: str, style: str | None) -> int | None:
        """Return the chapter number for a standalone 'Chapter N' marker line."""

        normalized_text = self._normalize_text(text)
        if not normalized_text:
            return None
        if self._looks_like_chapter_style(style):
            return None

        for pattern in self.chapter_patterns:
            match = pattern.match(normalized_text)
            if not match:
                continue
            title = self._normalize_text(match.groupdict().get("title", ""))
            if title:
                continue
            return self._coerce_chapter_number(match.group("number"))
        return None

    def _collect_paragraphs(self, doc: DocxDocument, limit: int | None = None) -> list[_ParagraphInfo]:
        """Return normalized paragraph data for the document."""

        collected: list[_ParagraphInfo] = []
        source = doc.paragraphs if limit is None else doc.paragraphs[:limit]

        for index, paragraph in enumerate(source):
            text = self._normalize_text(self._paragraph_text(paragraph))
            if not text:
                continue
            collected.append(
                _ParagraphInfo(
                    index=index,
                    text=text,
                    style=self._paragraph_style(paragraph),
                    paragraph=paragraph,
                )
            )

        return collected

    def _paragraph_text(self, paragraph: Paragraph) -> str:
        """Return paragraph text, recovering dropped-cap initials from floating textboxes."""

        text = paragraph.text or ""
        if not text:
            return text

        if not any(element.tag == qn("w:drawing") for element in paragraph._p.iter()):
            return text

        text_nodes = [element.text for element in paragraph._p.iter() if element.tag == qn("w:t") and element.text]
        if len(text_nodes) < 3:
            return text

        first = text_nodes[0].strip()
        second = text_nodes[1].strip()
        if (
            len(first) == 1
            and first.isalpha()
            and second.casefold() == first.casefold()
            and not text.startswith(first)
            and not text.startswith(first.lower())
        ):
            if first.upper() in {"A", "I", "O"} and self._drop_cap_prefix_wants_space(first, text):
                return f"{first} {text}"
            return first + text

        return text

    @staticmethod
    @lru_cache(maxsize=1)
    def _system_dictionary_words() -> frozenset[str]:
        """Return a best-effort system word list for drop-cap spacing heuristics."""

        dictionary_path = Path("/usr/share/dict/words")
        if not dictionary_path.exists():
            return frozenset()
        try:
            return frozenset(
                line.strip().casefold()
                for line in dictionary_path.read_text(encoding="utf-8", errors="ignore").splitlines()
                if line.strip()
            )
        except OSError:
            return frozenset()

    def _drop_cap_prefix_wants_space(self, prefix: str, text: str) -> bool:
        """Return whether a recovered drop-cap letter should stand as its own word."""

        if not text[:1].islower():
            return False

        first_word_match = re.match(r"[A-Za-z’']+", text)
        if first_word_match is None:
            return False

        combined_word = (prefix + first_word_match.group(0)).replace("’", "'").casefold()
        dictionary_words = self._system_dictionary_words()
        if not dictionary_words:
            return False
        return combined_word not in dictionary_words

    def _find_title(self, paragraphs: list[_ParagraphInfo]) -> tuple[int, str]:
        """Find the most likely title paragraph from the front matter."""

        for position, paragraph in enumerate(paragraphs[:6]):
            if paragraph.style and paragraph.style.lower() == "title" and not self._is_skip_section(paragraph.text):
                return position, paragraph.text

        for position, paragraph in enumerate(paragraphs[:6]):
            if self._is_skip_section(paragraph.text) or self._looks_like_credit_or_note(paragraph.text):
                continue
            if self._paragraph_is_emphasized(paragraph):
                return position, paragraph.text

        for position, paragraph in enumerate(paragraphs[:6]):
            if self._is_skip_section(paragraph.text) or self._looks_like_credit_or_note(paragraph.text):
                continue
            return position, paragraph.text

        raise ValueError("Unable to determine the book title from the opening paragraphs.")

    def _find_author(self, paragraphs: list[_ParagraphInfo], title_position: int) -> tuple[int, str]:
        """Find the author line in the front matter."""

        search_window = paragraphs[title_position + 1 : title_position + 10]
        for offset, paragraph in enumerate(search_window, start=title_position + 1):
            for pattern in self.author_patterns:
                match = pattern.match(paragraph.text)
                if match:
                    author_text = self._normalize_text(match.group("author"))
                    if author_text.casefold() in self.NON_AUTHOR_PHRASES:
                        logger.debug("Rejecting 'by' match as non-author phrase: %s", author_text)
                        continue
                    return offset, author_text

        for offset, paragraph in enumerate(search_window, start=title_position + 1):
            if self._looks_like_author_name(paragraph.text):
                return offset, paragraph.text

        logger.warning("Could not determine author from front matter; using 'Unknown Author'.")
        return title_position, "Unknown Author"

    def _find_subtitle(
        self,
        paragraphs: list[_ParagraphInfo],
        title_position: int,
        author_position: int,
    ) -> str | None:
        """Return subtitle text found between the title and author lines."""

        subtitle_parts: list[str] = []
        for paragraph in paragraphs[title_position + 1 : author_position]:
            if self._is_skip_section(paragraph.text):
                continue
            if self._looks_like_credit_or_note(paragraph.text):
                continue
            subtitle_parts.append(paragraph.text)

        subtitle = " ".join(subtitle_parts).strip()
        return subtitle or None

    def _find_original_publisher(self, paragraphs: list[_ParagraphInfo]) -> str | None:
        """Look for publisher metadata near the front of the document."""

        for paragraph in paragraphs[:20]:
            for pattern in self.publisher_patterns:
                match = pattern.match(paragraph.text)
                if match:
                    return self._normalize_text(match.group("publisher"))
        return None

    def _extract_author_from_folder(self, folder_name: str) -> str | None:
        """Attempt to extract an author name from a manuscript folder name."""

        cleaned = re.sub(r"^\d+[\s\-]+", "", folder_name)
        cleaned = re.sub(r"-EN-v-\d+.*$", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"-?\d+\.?\d*x\d+\.?\d*[-\s]*\d*\s*\d*$", "", cleaned)
        cleaned = cleaned.replace("-", " ").strip()
        cleaned_lower = cleaned.casefold()

        best_match: str | None = None
        best_length = 0
        for key, canonical_name in self.KNOWN_AUTHORS.items():
            if key in cleaned_lower and len(key) > best_length:
                best_match = canonical_name
                best_length = len(key)

        return best_match

    def _build_chapter(self, heading: dict[str, Any], body_paragraphs: list[str]) -> Chapter | None:
        """Build a chapter object from a heading plus collected body paragraphs."""

        raw_text = "\n\n".join(paragraph for paragraph in body_paragraphs if paragraph).strip()
        if not raw_text:
            logger.warning(
                "Skipping %s '%s' — heading found but no body text.",
                heading["type"],
                heading["title"],
            )
            return None

        return Chapter(
            number=heading["number"],
            title=heading["title"],
            type=heading["type"],
            raw_text=raw_text,
            word_count=self._count_words(raw_text),
        )

    def _is_generic_chapter_title(self, title: str) -> bool:
        """Return whether a chapter title is still just a synthetic numbered placeholder."""

        normalized_title = self._normalize_text(title)
        return bool(re.fullmatch(r"(chapter|book)\s+\d+", normalized_title, re.IGNORECASE))

    def _looks_like_followup_chapter_title(self, text: str, *, allow_long: bool) -> bool:
        """Return whether a body paragraph is actually the chapter's real title."""

        normalized_text = self._normalize_text(text)
        if not normalized_text:
            return False
        if self._is_skip_section(normalized_text) or self._is_back_matter_section(normalized_text):
            return False
        if self._looks_like_credit_or_note(normalized_text):
            return False
        if self._looks_like_author_heading_metadata(normalized_text):
            return False
        if re.match(r"^[\"'“”‘’(\[]", normalized_text):
            return False
        if normalized_text.endswith((":", ";")):
            return False

        words = normalized_text.split()
        max_words = 32 if allow_long else 5
        if not 1 <= len(words) <= max_words:
            return False
        if len(words) == 1 and re.search(r"[.?!]$", normalized_text):
            return False

        title_stop_words = {
            "a",
            "an",
            "and",
            "as",
            "at",
            "but",
            "by",
            "for",
            "from",
            "in",
            "into",
            "nor",
            "of",
            "on",
            "or",
            "the",
            "to",
            "with",
        }
        title_like_words = 0
        alpha_words = 0
        for word in words:
            for segment in re.split(r"[—–-]+", word):
                stripped = segment.strip(".,!?;:()[]{}\"'“”‘’—–-")
                if not stripped:
                    continue
                if not any(character.isalpha() for character in stripped):
                    continue
                alpha_words += 1
                if (
                    stripped.casefold() in title_stop_words
                    or stripped.isupper()
                    or re.fullmatch(r"[ivxlcdm]+", stripped, re.IGNORECASE)
                    or (stripped[:1].isupper() and stripped[1:] == stripped[1:].lower())
                ):
                    title_like_words += 1

        if alpha_words == 0:
            return False

        allowed_non_title_words = 1 if allow_long and alpha_words >= 6 else 0
        if title_like_words < alpha_words - allowed_non_title_words:
            return False

        return not re.search(r"\b(?:said|asked|replied|cried|answered|whispered)\b[.?!]?$", normalized_text, re.IGNORECASE)

    def _validate_toc(self, chapters: list[Chapter]) -> None:
        """Log TOC mismatches when a table of contents was detected."""

        if not self.last_toc_entries:
            return

        normalized_toc = {self._comparison_key(entry) for entry in self.last_toc_entries}
        chapter_titles = {
            self._comparison_key(chapter.title)
            for chapter in chapters
            if chapter.type == "chapter"
        }
        missing = sorted(title for title in normalized_toc if title and title not in chapter_titles)
        if missing:
            logger.debug("TOC entries without matching detected chapters: %s", missing)

    def _looks_like_author_name(self, text: str) -> bool:
        """Return whether text looks like a plain author name."""

        normalized_text = self._normalize_text(text)
        if normalized_text.casefold() in self.NON_AUTHOR_PHRASES:
            return False
        if re.match(
            r"^(a |the )?(modern|new|contemporary|fresh|complete|selected|collected|essential)\b",
            normalized_text,
            re.IGNORECASE,
        ):
            return False
        if self._is_chapter_heading(normalized_text, None)[0]:
            return False
        if self._looks_like_credit_or_note(normalized_text):
            return False
        if re.search(r"\d", normalized_text):
            return False

        words = normalized_text.split()
        if not 2 <= len(words) <= 6:
            return False

        return all(re.match(r"^[A-Z][A-Za-z'.’-]*$", word.strip(",.;:")) for word in words)

    def _looks_like_author_heading_metadata(self, text: str) -> bool:
        """Return whether a heading is likely just an author-credit line."""

        normalized_text = self._normalize_text(text).casefold()
        if not normalized_text or len(normalized_text.split()) > 3:
            return False
        return normalized_text in self.KNOWN_AUTHORS

    def _looks_like_generic_section_divider(self, text: str) -> bool:
        """Return whether a heading is just a structural divider, not narratable content."""

        normalized_text = self._normalize_heading_for_skip_rules(text)
        return bool(
            re.match(
                r"^(part|section|volume)\s+(?:\d+|[ivxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)$",
                normalized_text,
                re.IGNORECASE,
            )
            or normalized_text in {"shang pian", "xia pian", "upper part", "lower part"}
        )

    def _is_running_header_paragraph(
        self,
        text: str,
        style: str | None,
        running_headers: set[str],
        *,
        metadata_title: str | None = None,
        current_heading_title: str | None = None,
    ) -> bool:
        """Return whether a paragraph line is likely a PDF-conversion running header."""

        if not style or style.casefold() != "normal":
            return False

        normalized_text = self._normalize_text(text)
        if normalized_text.casefold() in running_headers:
            return True

        normalized_title = self._normalize_text(metadata_title).casefold() if metadata_title else ""
        if normalized_title and (
            normalized_text.casefold() == normalized_title
            or normalized_text.casefold().startswith(f"{normalized_title}:")
            or normalized_text.casefold().startswith(f"{normalized_title} ")
        ):
            return True

        if current_heading_title and ":" in normalized_text:
            _prefix, suffix = normalized_text.split(":", 1)
            if self._comparison_key(suffix) == self._comparison_key(current_heading_title):
                return True

        return False

    def _looks_like_credit_or_note(self, text: str) -> bool:
        """Return whether a line is clearly metadata but not subtitle content."""

        normalized_text = self._normalize_text(text)
        return bool(
            re.match(r"^(translated|adapted|edited|illustrated)\s+by\b", normalized_text, re.IGNORECASE)
            or re.match(r"^(visit|www\.|https?://)", normalized_text, re.IGNORECASE)
            or "libraryofalexandria.com" in normalized_text.lower()
            or re.match(
                r"^(a |the )?(modern|new|contemporary|fresh|complete|selected|collected|essential)\s+"
                r"(translation|edition|works|writings)\b",
                normalized_text,
                re.IGNORECASE,
            )
            or re.match(r"^(volume|part)\s+[ivxlcdm\d]+\b", normalized_text, re.IGNORECASE)
            or re.match(
                r"^(first|second|third|revised|annotated|unabridged|definitive)\s+edition\b",
                normalized_text,
                re.IGNORECASE,
            )
        )

    def _looks_like_toc_entry(self, text: str, style: str | None) -> bool:
        """Return whether a paragraph looks like a TOC entry."""

        if self._looks_like_chapter_style(style):
            return False
        if self._looks_like_credit_or_note(text):
            return False
        if style and "toc" in style.lower():
            return True
        if re.search(r"(?:\t|\s)\d+\s*$", text):
            return True
        return len(text.split()) <= 16

    def _is_toc_heading(self, text: str) -> bool:
        """Return whether text marks the start of a table of contents."""

        return bool(re.match(r"^(table of contents|contents)\b", self._normalize_text(text), re.IGNORECASE))

    def _is_back_matter_section(self, text: str) -> bool:
        """Return whether text marks terminal back matter."""

        normalized_text = self._normalize_heading_for_skip_rules(text)
        return any(pattern.match(normalized_text) for pattern in self.back_matter_patterns)

    def _looks_like_chapter_style(self, style: str | None) -> bool:
        """Return whether a paragraph style looks like a chapter heading style."""

        if not style:
            return False
        style_lower = style.lower()
        return style_lower.startswith("heading")

    def _looks_like_primary_section_style(self, style: str | None) -> bool:
        """Return whether a paragraph style looks like a top-level section heading."""

        if not style:
            return False
        return bool(re.match(r"^heading\s*1\b", style, re.IGNORECASE))

    def _heading_level(self, style: str | None) -> int | None:
        """Return the numeric heading level encoded in a DOCX style name."""

        if not style:
            return None
        match = re.match(r"^heading\s*(\d+)\b", style, re.IGNORECASE)
        if match is None:
            return None
        return int(match.group(1))

    def _detect_primary_numbered_heading_level(self, paragraphs: list[_ParagraphInfo]) -> int | None:
        """Infer the dominant heading level used for numbered narrative sections."""

        numbered_levels: list[int] = []
        keyword_level_counts: Counter[int] = Counter()
        first_keyword_index = next(
            (
                paragraph.index
                for paragraph in paragraphs
                if self._heading_level(paragraph.style) is not None
                and self._matches_keyword_numbered_section_heading(paragraph.text)
            ),
            None,
        )
        for paragraph in paragraphs:
            heading_level = self._heading_level(paragraph.style)
            if heading_level is None:
                continue
            is_keyword_heading = self._matches_keyword_numbered_section_heading(paragraph.text)
            if (
                first_keyword_index is not None
                and paragraph.index < first_keyword_index
                and not is_keyword_heading
            ):
                continue
            if self._matches_numbered_section_heading(paragraph.text):
                numbered_levels.append(heading_level)
                if is_keyword_heading:
                    keyword_level_counts[heading_level] += 1

        if not numbered_levels:
            return None

        level_counts = Counter(numbered_levels)
        repeating_levels = [level for level, count in level_counts.items() if count >= 2]
        if repeating_levels:
            return max(
                repeating_levels,
                key=lambda level: (level_counts[level], keyword_level_counts[level], level),
            )
        return min(numbered_levels)

    def _detect_dominant_heading_level(self, paragraphs: list[_ParagraphInfo]) -> int | None:
        """
        Return the most common non-skippable heading level when no numbered headings exist.

        The threshold avoids promoting a small handful of metadata headings into chapters.
        """

        level_counts: dict[int, int] = {}
        for paragraph in paragraphs:
            heading_level = self._heading_level(paragraph.style)
            if heading_level is None:
                continue
            text = self._normalize_text(paragraph.text)
            if not text or self._is_skip_section(text):
                continue
            if self._is_back_matter_section(text):
                continue
            level_counts[heading_level] = level_counts.get(heading_level, 0) + 1

        if not level_counts:
            return None

        max_count = max(level_counts.values())
        if max_count < 15:
            return None

        return max(level_counts, key=lambda level: (level_counts[level], level))

    def _matches_numbered_section_heading(self, text: str) -> bool:
        """Return whether a heading looks like a numbered chapter/book boundary."""

        normalized_text = self._normalize_text(text)
        if not normalized_text or self._is_skip_section(normalized_text):
            return False
        return any(
            pattern.match(normalized_text)
            for pattern in self.chapter_patterns + self.book_patterns + self.part_patterns + self.maxim_patterns
        )

    def _matches_keyword_numbered_section_heading(self, text: str) -> bool:
        """Return whether text uses an explicit section keyword, not just a bare numeral."""

        normalized_text = self._normalize_text(text)
        if not normalized_text or self._is_skip_section(normalized_text):
            return False
        return any(
            pattern.match(normalized_text)
            for pattern in (self.chapter_patterns[0],) + self.book_patterns + self.part_patterns + self.maxim_patterns
        )

    def _paragraph_is_emphasized(self, paragraph: _ParagraphInfo) -> bool:
        """Return whether paragraph formatting makes it a strong title candidate."""

        if paragraph.style and "title" in paragraph.style.lower():
            return True
        return any(run.bold for run in paragraph.paragraph.runs)

    def _paragraph_style(self, paragraph: Paragraph) -> str | None:
        """Return the paragraph style name when available."""

        return paragraph.style.name if paragraph.style is not None else None

    def _normalize_text(self, text: str) -> str:
        """Collapse internal whitespace while preserving characters."""

        return " ".join(text.replace("\xa0", " ").split())

    def _normalize_heading_for_skip_rules(self, text: str) -> str:
        """Normalize heading text for punctuation-insensitive skip comparisons."""

        normalized = self._normalize_text(text).casefold()
        normalized = re.sub(r"[^\w\s]", " ", normalized)
        return re.sub(r"\s+", " ", normalized).strip()

    def _comparison_key(self, text: str) -> str:
        """Normalize heading text for loose comparisons."""

        return re.sub(r"[^a-z0-9]+", "", text.casefold())

    def _coerce_chapter_number(self, value: str) -> int | None:
        """Convert a chapter number token into an integer when possible."""

        normalized_value = value.strip().strip(".:").replace("-", " ").casefold()
        if normalized_value.isdigit():
            return int(normalized_value)
        if re.fullmatch(r"[ivxlcdm]+", normalized_value, re.IGNORECASE):
            return self._roman_to_int(normalized_value.upper())
        if normalized_value in self.word_number_map:
            return self.word_number_map[normalized_value]
        return None

    def _roman_to_int(self, value: str) -> int:
        """Convert a Roman numeral string into an integer."""

        numerals = {
            "I": 1,
            "V": 5,
            "X": 10,
            "L": 50,
            "C": 100,
            "D": 500,
            "M": 1000,
        }
        total = 0
        previous = 0

        for character in reversed(value):
            current = numerals[character]
            if current < previous:
                total -= current
            else:
                total += current
                previous = current

        return total
