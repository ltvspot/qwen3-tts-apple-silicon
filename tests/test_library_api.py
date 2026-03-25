"""Library scanner and parser API tests."""

from __future__ import annotations

import threading
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from src.api.library import LibraryScanner
from src.config import settings
from src.database import Book, Chapter
from tests.parser_fixture_utils import create_sample_epub, create_sample_pdf, default_book_sections


def _create_sample_docx(docx_path: Path) -> None:
    """Create a synthetic manuscript that exercises intro and chapter parsing."""

    document = Document()
    document.add_paragraph("The Test Chronicle", style="Title")
    document.add_paragraph("A Detective Story")
    document.add_paragraph("by Jane Doe")
    document.add_paragraph("Copyright 2026")
    document.add_paragraph("Table of Contents")
    document.add_paragraph("Introduction")
    document.add_paragraph("Chapter I. The Beginning")
    document.add_paragraph("Introduction", style="Heading 1")
    document.add_paragraph("Opening setup.")
    document.add_paragraph("Chapter I. The Beginning", style="Heading 1")
    document.add_paragraph("The story starts here.")
    document.add_paragraph("Thank You for Reading")
    document.save(docx_path)


def _create_library_folder(root: Path, folder_name: str, *, with_docx: bool = True) -> Path:
    """Create a temporary manuscript folder for tests."""

    folder = root / folder_name
    folder.mkdir(parents=True, exist_ok=True)
    if with_docx:
        _create_sample_docx(folder / f"{folder_name}-Word-6x9-Clean.docx")
    return folder


def _create_epub_library_folder(root: Path, folder_name: str) -> Path:
    """Create an EPUB-only manuscript folder for parser fallback tests."""

    folder = root / folder_name
    folder.mkdir(parents=True, exist_ok=True)
    create_sample_epub(
        folder / f"{folder_name}.epub",
        title="The EPUB Chronicle",
        subtitle="A Digital Mystery",
        author="Jane Doe",
        sections=default_book_sections(),
    )
    return folder


def _create_pdf_library_folder(root: Path, folder_name: str) -> Path:
    """Create a PDF-only manuscript folder for parser fallback tests."""

    folder = root / folder_name
    folder.mkdir(parents=True, exist_ok=True)
    create_sample_pdf(
        folder / f"{folder_name}.pdf",
        title="The PDF Chronicle",
        subtitle="A Portable Mystery",
        author="Jane Doe",
        sections=default_book_sections(),
    )
    return folder


def test_library_scanner_parses_realistic_folder_variants() -> None:
    """Scanner metadata parsing should tolerate the existing folder naming drift."""

    scanner = LibraryScanner()

    standard = scanner._parse_folder_name("0936.-The-Murder-of-Roger-Ackroyd-6x9-234")
    epub_only = scanner._parse_folder_name("204-The-Ultimate-Horror-Collection-EPub")
    malformed = scanner._parse_folder_name("drive-download-20251111T072835Z-1-001")

    assert standard == {
        "id": "0936",
        "title": "The Murder of Roger Ackroyd",
        "trim_size": "6x9",
        "page_count": 234,
    }
    assert epub_only == {
        "id": "204",
        "title": "The Ultimate Horror Collection",
        "trim_size": None,
        "page_count": None,
    }
    assert malformed is None


def test_scan_library_and_get_library(
    client: TestClient,
    test_db: Session,
    tmp_path: Path,
    monkeypatch: object,
) -> None:
    """Scan folders into the DB and expose them through the library endpoints."""

    _create_library_folder(tmp_path, "1001-The-Hidden-Library-6x9-123")
    epub_folder = tmp_path / "1002-The-EPub-Only-Collection-EPub"
    epub_folder.mkdir()
    (epub_folder / "1002-The-EPub-Only-Collection.epub").write_text("stub", encoding="utf-8")
    (tmp_path / "drive-download-20251111T072835Z-1-001").mkdir()

    monkeypatch.setattr(settings, "FORMATTED_MANUSCRIPTS_PATH", str(tmp_path))

    response = client.post("/api/library/scan")
    assert response.status_code == 200
    assert response.json() == {
        "total_found": 3,
        "total_indexed": 2,
        "new_books": 2,
        "errors": ["Unable to parse folder metadata: drive-download-20251111T072835Z-1-001"],
    }

    stored_books = test_db.query(Book).order_by(Book.folder_path).all()
    assert [book.folder_path for book in stored_books] == [
        "1001-The-Hidden-Library-6x9-123",
        "1002-The-EPub-Only-Collection-EPub",
    ]
    assert stored_books[0].title == "The Hidden Library"
    assert stored_books[0].author == "Unknown Author"
    assert stored_books[0].status == "not_started"
    assert stored_books[1].trim_size is None

    library_response = client.get("/api/library")
    assert library_response.status_code == 200
    assert library_response.json() == {
        "total": 2,
        "books": [
            {
                "id": stored_books[0].id,
                "title": "The Hidden Library",
                "subtitle": None,
                "author": "Unknown Author",
                "narrator": "Kent Zimering",
                "folder_path": "1001-The-Hidden-Library-6x9-123",
                "status": "not_started",
                "page_count": 123,
                "trim_size": "6x9",
                "chapter_count": 0,
                "created_at": stored_books[0].created_at.isoformat(),
                "updated_at": stored_books[0].updated_at.isoformat(),
                "generation_status": "idle",
                "generation_started_at": None,
                "generation_eta_seconds": None,
            },
            {
                "id": stored_books[1].id,
                "title": "The EPub Only Collection",
                "subtitle": None,
                "author": "Unknown Author",
                "narrator": "Kent Zimering",
                "folder_path": "1002-The-EPub-Only-Collection-EPub",
                "status": "not_started",
                "page_count": None,
                "trim_size": None,
                "chapter_count": 0,
                "created_at": stored_books[1].created_at.isoformat(),
                "updated_at": stored_books[1].updated_at.isoformat(),
                "generation_status": "idle",
                "generation_started_at": None,
                "generation_eta_seconds": None,
            },
        ],
        "stats": {
            "not_started": 2,
            "parsed": 0,
            "generating": 0,
            "generated": 0,
            "qa": 0,
            "qa_approved": 0,
            "exported": 0,
        },
    }


def test_scan_progress_endpoint_reports_live_and_completed_state(
    client: TestClient,
    monkeypatch: object,
) -> None:
    """Scan progress should be visible while a scan is running and after it completes."""

    progress_started = threading.Event()
    release_scan = threading.Event()

    def fake_scan(self, db_session, *, progress_callback=None):  # noqa: ANN001
        del self, db_session
        if progress_callback is not None:
            progress_callback({
                "errors": [],
                "files_found": 10,
                "files_processed": 3,
                "new_books": 1,
            })
        progress_started.set()
        assert release_scan.wait(timeout=2)
        if progress_callback is not None:
            progress_callback({
                "errors": [],
                "files_found": 10,
                "files_processed": 10,
                "new_books": 2,
            })
        return {
            "errors": [],
            "new_books": 2,
            "total_found": 10,
            "total_indexed": 10,
        }

    monkeypatch.setattr(LibraryScanner, "scan", fake_scan)

    response_holder: dict[str, object] = {}

    def run_scan() -> None:
        response_holder["response"] = client.post("/api/library/scan")

    thread = threading.Thread(target=run_scan)
    thread.start()

    assert progress_started.wait(timeout=2)
    in_flight_response = client.get("/api/library/scan/progress")
    assert in_flight_response.status_code == 200
    assert in_flight_response.json()["scanning"] is True
    assert in_flight_response.json()["files_found"] == 10
    assert in_flight_response.json()["files_processed"] == 3

    release_scan.set()
    thread.join(timeout=2)

    completed_response = response_holder["response"]
    assert completed_response.status_code == 200

    final_progress = client.get("/api/library/scan/progress")
    assert final_progress.status_code == 200
    assert final_progress.json()["scanning"] is False
    assert final_progress.json()["files_found"] == 10
    assert final_progress.json()["files_processed"] == 10
    assert final_progress.json()["new_books"] == 2


def test_parse_book_flow_and_chapter_updates(
    client: TestClient,
    test_db: Session,
    tmp_path: Path,
    monkeypatch: object,
) -> None:
    """Parse an indexed DOCX manuscript, expose chapters, and support overwrite-safe edits."""

    folder_name = "1003-Placeholder-Title-6x9-145"
    _create_library_folder(tmp_path, folder_name)
    monkeypatch.setattr(settings, "FORMATTED_MANUSCRIPTS_PATH", str(tmp_path))

    scan_response = client.post("/api/library/scan")
    assert scan_response.status_code == 200

    book = test_db.query(Book).filter(Book.folder_path == folder_name).one()

    parsed_before = client.get(f"/api/book/{book.id}/parsed")
    assert parsed_before.status_code == 400
    assert parsed_before.json() == {"detail": "Book not yet parsed. Status: BookStatus.NOT_STARTED"}

    parse_response = client.post(f"/api/book/{book.id}/parse", json={})
    assert parse_response.status_code == 200
    assert parse_response.json() == {
        "status": "parsing",
        "chapters_detected": 4,
        "message": "Successfully parsed 2 narratable sections plus opening/closing credits.",
    }

    test_db.refresh(book)
    assert book.title == "The Test Chronicle"
    assert book.subtitle == "A Detective Story"
    assert book.author == "Jane Doe"
    assert book.status == "parsed"

    chapters = test_db.query(Chapter).filter(Chapter.book_id == book.id).order_by(Chapter.number).all()
    assert [chapter.number for chapter in chapters] == [0, 1, 2, 3]
    assert [chapter.type.value for chapter in chapters] == [
        "opening_credits",
        "introduction",
        "chapter",
        "closing_credits",
    ]

    parsed_response = client.get(f"/api/book/{book.id}/parsed")
    assert parsed_response.status_code == 200
    parsed_payload = parsed_response.json()
    assert [chapter["number"] for chapter in parsed_payload] == [0, 1, 2, 3]
    assert parsed_payload[1]["title"] == "Introduction"
    assert parsed_payload[2]["title"] == "The Beginning"

    already_parsed = client.post(f"/api/book/{book.id}/parse", json={"overwrite": False})
    assert already_parsed.status_code == 200
    assert already_parsed.json() == {
        "status": "already_parsed",
        "chapters_detected": 4,
        "message": "Book already parsed. Set overwrite=True to re-parse.",
    }

    update_response = client.put(
        f"/api/book/{book.id}/chapter/2/text",
        json={"text_content": "The revised chapter text now has five words."},
    )
    assert update_response.status_code == 200
    assert update_response.json()["word_count"] == 8
    assert update_response.json()["text_content"] == "The revised chapter text now has five words."

    blank_update_response = client.put(
        f"/api/book/{book.id}/chapter/2/text",
        json={"text_content": "   "},
    )
    assert blank_update_response.status_code == 400
    assert blank_update_response.json() == {"detail": "Chapter text cannot be empty."}

    overwrite_response = client.post(f"/api/book/{book.id}/parse", json={"overwrite": True})
    assert overwrite_response.status_code == 200
    assert overwrite_response.json()["chapters_detected"] == 4
    assert test_db.query(Chapter).filter(Chapter.book_id == book.id).count() == 4


def test_parse_book_uses_epub_fallback(
    client: TestClient,
    test_db: Session,
    tmp_path: Path,
    monkeypatch: object,
) -> None:
    """Parsing should succeed with EPUB when a DOCX manuscript is unavailable."""

    folder_name = "1004-EPub-Only-Book-EPub"
    _create_epub_library_folder(tmp_path, folder_name)
    monkeypatch.setattr(settings, "FORMATTED_MANUSCRIPTS_PATH", str(tmp_path))

    scan_response = client.post("/api/library/scan")
    assert scan_response.status_code == 200
    book = test_db.query(Book).filter(Book.folder_path == folder_name).one()

    parse_response = client.post(f"/api/book/{book.id}/parse", json={})
    assert parse_response.status_code == 200
    assert parse_response.json()["chapters_detected"] == 5

    test_db.refresh(book)
    assert book.title == "The EPUB Chronicle"
    assert book.subtitle == "A Digital Mystery"
    assert book.author == "Jane Doe"


def test_parse_book_uses_pdf_fallback(
    client: TestClient,
    test_db: Session,
    tmp_path: Path,
    monkeypatch: object,
) -> None:
    """Parsing should succeed with PDF when DOCX and EPUB manuscripts are unavailable."""

    folder_name = "1005-PDF-Only-Book-PDF"
    _create_pdf_library_folder(tmp_path, folder_name)
    monkeypatch.setattr(settings, "FORMATTED_MANUSCRIPTS_PATH", str(tmp_path))

    scan_response = client.post("/api/library/scan")
    assert scan_response.status_code == 200
    book = test_db.query(Book).filter(Book.folder_path == folder_name).one()

    parse_response = client.post(f"/api/book/{book.id}/parse", json={})
    assert parse_response.status_code == 200
    assert parse_response.json()["chapters_detected"] == 5

    test_db.refresh(book)
    assert book.title == "The PDF Chronicle"
    assert book.subtitle == "A Portable Mystery"
    assert book.author == "Jane Doe"


def test_parse_book_requires_supported_format(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: object,
) -> None:
    """Parsing should fail cleanly when a folder has no supported manuscript format."""

    folder = tmp_path / "1006-Unsupported-Book-6x9-101"
    folder.mkdir()
    (folder / "notes.txt").write_text("unsupported", encoding="utf-8")
    monkeypatch.setattr(settings, "FORMATTED_MANUSCRIPTS_PATH", str(tmp_path))

    scan_response = client.post("/api/library/scan")
    assert scan_response.status_code == 200
    book_id = client.get("/api/library").json()["books"][0]["id"]

    parse_response = client.post(f"/api/book/{book_id}/parse", json={})
    assert parse_response.status_code == 400
    assert parse_response.json() == {"detail": "No supported manuscript format in 1006-Unsupported-Book-6x9-101"}
