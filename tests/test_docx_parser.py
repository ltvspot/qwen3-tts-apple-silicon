"""DOCX parser and text cleaning tests."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from src.parser import CreditsGenerator, DocxParser, ManuscriptParserFactory, TextCleaner
from src.parser.common import should_skip_heading


def _find_sherlock_docx() -> Path | None:
    """Return the preferred Sherlock Holmes DOCX test manuscript."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("0906*/*Clean.docx"))
    if candidates:
        return candidates[0]

    fallback_candidates = sorted(manuscripts_path.glob("0906*/*.docx"))
    return fallback_candidates[0] if fallback_candidates else None


def _find_meditations_docx() -> Path | None:
    """Return the Meditations DOCX used to guard hierarchical heading parsing."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("002-Meditations*/*.docx"))
    return candidates[0] if candidates else None


def _find_tao_te_ching_docx() -> Path | None:
    """Return the Tao Te Ching DOCX used to guard against false section splits."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("003-Tao-Te-Ching*/*.docx"))
    return candidates[0] if candidates else None


def _find_as_a_man_thinketh_folder() -> Path | None:
    """Return the As a Man Thinketh manuscript folder used for hierarchy regression tests."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("005-As-a-Man-Thinketh*"))
    return candidates[0] if candidates else None


def _find_book_of_five_rings_folder() -> Path | None:
    """Return the Book of Five Rings manuscript folder for heading-title regressions."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("004-The-Book-of-Five-Rings*"))
    return candidates[0] if candidates else None


def _find_beyond_good_and_evil_folder() -> Path | None:
    """Return the Beyond Good and Evil manuscript folder for heading-title regressions."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("010-Beyond-Good-and-Evil*"))
    return candidates[0] if candidates else None


def _find_nicomachean_ethics_folder() -> Path | None:
    """Return the Nicomachean Ethics manuscript folder for book-sequence regressions."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("011-Nicomachean-Ethics*"))
    return candidates[0] if candidates else None


def _find_oliver_twist_folder() -> Path | None:
    """Return the Oliver Twist manuscript folder for follow-up title regressions."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("0829-Oliver-Twist*"))
    return candidates[0] if candidates else None


def _find_bhagavad_gita_folder() -> Path | None:
    """Return the Bhagavad Gita manuscript folder for speaker-label regressions."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("1039-The-Bhagavad-Gita*"))
    return candidates[0] if candidates else None


def _find_shortness_of_life_folder() -> Path | None:
    """Return the On the Shortness of Life manuscript folder for front-matter regression tests."""

    manuscripts_path = Path(__file__).resolve().parent.parent / "Formatted Manuscripts"
    candidates = sorted(manuscripts_path.glob("007-On-the-Shortness-of-Life*"))
    return candidates[0] if candidates else None


def _write_docx(docx_path: Path, paragraphs: list[tuple[str, str | None]]) -> None:
    """Create a DOCX file from `(text, style)` paragraph pairs."""

    document = Document()
    available_styles = {style.name for style in document.styles}

    for _text, style in paragraphs:
        if style is not None and style not in available_styles:
            document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
            available_styles.add(style)

    for text, style in paragraphs:
        if style is None:
            document.add_paragraph(text)
        else:
            document.add_paragraph(text, style=style)
    document.save(docx_path)


def test_parse_sherlock_holmes_manuscript() -> None:
    """Parse the real Sherlock Holmes manuscript and validate core structure."""

    parser = DocxParser()
    test_docx = _find_sherlock_docx()

    assert test_docx is not None, "Sherlock Holmes manuscript not found."

    metadata, chapters = parser.parse(test_docx)

    assert metadata.title == "The Adventures of Sherlock Holmes"
    assert metadata.author == "Arthur Conan Doyle"
    assert metadata.subtitle is not None
    assert "Mystery" in metadata.subtitle

    assert len(chapters) == 13
    assert chapters[0].number == 0
    assert chapters[0].type == "introduction"
    assert chapters[0].title == "Introduction"
    assert chapters[0].word_count > 100
    assert "\n\n" in chapters[0].raw_text

    numbered_chapters = [chapter for chapter in chapters if chapter.type == "chapter"]
    assert len(numbered_chapters) == 12
    assert numbered_chapters[0].number == 1
    assert numbered_chapters[0].title == "A Scandal in Bohemia"
    assert numbered_chapters[-1].number == 12
    assert numbered_chapters[-1].title == "The Adventure of The Copper Beeches"
    assert all("Preface Message to the Reader" not in chapter.raw_text for chapter in chapters)
    assert all("Thank You for Reading" not in chapter.raw_text for chapter in chapters)


def test_parse_meditations_manuscript_respects_book_hierarchy() -> None:
    """Meditations should parse into one intro plus its twelve real books."""

    parser = DocxParser()
    test_docx = _find_meditations_docx()

    assert test_docx is not None, "Meditations manuscript not found."

    metadata, chapters = parser.parse(test_docx)

    assert metadata.author == "Marcus Aurelius"
    assert [chapter.title for chapter in chapters] == [
        "Introduction",
        "Book One",
        "Book Two",
        "Book Three",
        "Book Four",
        "Book Five",
        "Book Six",
        "Book Seven",
        "Book Eight",
        "Book Nine",
        "Book Ten",
        "Book Eleven",
        "Book Twelve",
    ]
    assert "The Meditations" in chapters[0].raw_text
    assert "Daily Reflection and Journaling" in chapters[0].raw_text
    assert "From Fronto, about how much envy" in chapters[1].raw_text
    assert "Chapter 16" in chapters[3].raw_text
    assert all(chapter.title != "Chapter 16" for chapter in chapters)
    assert all(chapter.title != "Daily Reflection and Journaling" for chapter in chapters)
    assert "Book One 35" in parser.last_toc_entries


def test_parse_tao_te_ching_manuscript_ignores_part_dividers() -> None:
    """Tao Te Ching should keep Shang Pian/Xia Pian as structure, not chapters."""

    parser = DocxParser()
    test_docx = _find_tao_te_ching_docx()

    assert test_docx is not None, "Tao Te Ching manuscript not found."

    metadata, chapters = parser.parse(test_docx)

    assert metadata.author == "Lao Tzu"
    assert len(chapters) == 82
    assert len([chapter for chapter in chapters if chapter.type == "chapter"]) == 81
    assert chapters[0].title == "Introduction"
    assert chapters[1].title == "Chapter 1"
    assert chapters[-1].title == "Chapter 81"
    assert all(chapter.title != "Shang Pian" for chapter in chapters)
    assert all(chapter.title != "Xia Pian" for chapter in chapters)
    assert all("Shang Pian" not in {line.strip() for line in chapter.raw_text.splitlines()} for chapter in chapters)
    assert all("Xia Pian" not in {line.strip() for line in chapter.raw_text.splitlines()} for chapter in chapters)
    assert all("Tao Te Ching" not in {line.strip() for line in chapter.raw_text.splitlines()} for chapter in chapters)
    assert "The End" not in chapters[-1].raw_text


def test_parse_book_of_five_rings_uses_real_heading_titles() -> None:
    """Book 4 should use H2 title headings as chapter titles and keep them out of body text."""

    folder = _find_book_of_five_rings_folder()
    assert folder is not None, "Book of Five Rings manuscript folder not found."

    metadata, chapters, manuscript_path = ManuscriptParserFactory.parse_manuscript(folder)

    assert metadata.title == "The Book of Five Rings"
    assert manuscript_path.name == "004 The Book of Five Rings-5x8-097.docx"
    assert [chapter.title for chapter in chapters] == [
        "Preface",
        "THE GROUND BOOK",
        "THE WATER BOOK",
        "THE FIRE BOOK",
        "THE WIND BOOK",
        "THE BOOK OF THE VOID",
    ]
    assert chapters[1].raw_text.startswith("Strategy is the skill of the warrior.")
    assert chapters[2].raw_text.startswith("The spirit of the Ni Ten Ichi school of strategy")
    assert "THE GROUND BOOK" not in {line.strip() for line in chapters[1].raw_text.splitlines()}
    assert "THE WATER BOOK" not in {line.strip() for line in chapters[2].raw_text.splitlines()}


def test_parse_beyond_good_and_evil_uses_real_heading_titles() -> None:
    """Book 11 should use H2 section titles, not synthetic Chapter N labels."""

    folder = _find_beyond_good_and_evil_folder()
    assert folder is not None, "Beyond Good and Evil manuscript folder not found."

    metadata, chapters, manuscript_path = ManuscriptParserFactory.parse_manuscript(folder)

    assert metadata.title == "Beyond Good and Evil"
    assert manuscript_path.name == "010-Beyond-Good-and-Evil-Friedrich-Nietzsche-6x9-262.docx"
    assert [chapter.title for chapter in chapters[:10]] == [
        "Introduction",
        "ON THE PREJUDICES OF PHILOSOPHERS",
        "THE FREE SPIRIT",
        "WHAT IS RELIGIOUS",
        "APOPHTHEGMS AND INTERLUDES",
        "THE NATURAL HISTORY OF MORALS",
        "WE SCHOLARS",
        "OUR VIRTUES",
        "PEOPLES AND FATHERLANDS",
        "WHAT IS NOBLE?",
    ]
    assert chapters[1].raw_text.startswith("The Will to Truth")
    assert chapters[6].raw_text.startswith("At the risk of sounding preachy")
    assert all(not chapter.title.startswith("Chapter ") for chapter in chapters[1:10])


def test_parse_oliver_twist_uses_following_plain_subtitles() -> None:
    """Chapter subtitles on plain paragraphs should become the chapter title."""

    folder = _find_oliver_twist_folder()
    assert folder is not None, "Oliver Twist manuscript folder not found."

    metadata, chapters, manuscript_path = ManuscriptParserFactory.parse_manuscript(folder)

    assert metadata.title == "Oliver Twist"
    assert manuscript_path.name == "0829-Oliver-Twist-Word-01-251020-6x9-Clean.docx"
    assert [chapter.title for chapter in chapters[:4]] == [
        "Introduction",
        "Describing the Place Where Oliver Twist Was Born and the Circumstances of His Birth",
        "Concerning Oliver Twist’s Growth, Education, And Upbringing",
        "Relates How Oliver Twist Came Very Close to Getting a Position That Would Have Been Anything but Easy",
    ]
    assert chapters[4].title == "Oliver is Offered Another Position and Begins His Public Life"
    assert (
        chapters[5].title
        == "Oliver Mingles with New Associates. Going to a Funeral for The First Time, He Forms an Unfavourable Notion of His Master’s Business"
    )
    assert (
        chapters[10].title
        == "Oliver Becomes More Familiar with The Characters of His New Companions and Gains Some Experience at a Great Cost. A Short, But Very Important Chapter in This History."
    )
    assert chapters[1].raw_text.startswith("Among the various public buildings in a certain town")
    assert chapters[2].raw_text.startswith("For the following eight to ten months")


def test_parse_bhagavad_gita_keeps_generic_chapter_titles_when_body_starts_with_speaker() -> None:
    """Speaker labels should stay in the body, not become the chapter title."""

    folder = _find_bhagavad_gita_folder()
    assert folder is not None, "Bhagavad Gita manuscript folder not found."

    metadata, chapters, manuscript_path = ManuscriptParserFactory.parse_manuscript(folder)

    assert metadata.title == "The Bhagavad Gita"
    assert manuscript_path.name == "1039-The-Bhagavad-Gita-Word-01-251020-5x8-Clean.docx"
    assert [chapter.title for chapter in chapters[:4]] == [
        "Introduction",
        "Chapter 1",
        "Chapter 2",
        "Chapter 3",
    ]
    assert chapters[1].raw_text.startswith("Dhritirashtra: On the sacred plain of Kurukshetra")
    assert chapters[2].raw_text.startswith("Sanjaya said:")
    assert chapters[3].raw_text.startswith("Arjuna.")


def test_running_headers_filtered_from_body(tmp_path: Path) -> None:
    """Repeated Normal-style running headers should not bleed into chapter text."""

    docx_path = tmp_path / "running-headers.docx"
    _write_docx(
        docx_path,
        [
            ("A Short Treatise", "Title"),
            ("by Jane Doe", None),
            ("Chapter 1", "Heading 2"),
            ("Actual body opening.", "Body Text"),
            ("Shang Pian", "Normal"),
            ("The wheel is useful because of the empty space in the center.", "Body Text"),
            ("Shang Pian", "Normal"),
            ("Mold clay into a pot.", "Body Text"),
            ("Shang Pian", "Normal"),
            ("Cut doors and windows to make a room.", "Body Text"),
            ("Shang Pian", "Normal"),
            ("Usefulness comes from what is not there.", "Body Text"),
            ("Shang Pian", "Normal"),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert len(chapters) == 1
    assert chapters[0].title == "Chapter 1"
    assert "Actual body opening." in chapters[0].raw_text
    assert "Usefulness comes from what is not there." in chapters[0].raw_text
    assert "Shang Pian" not in {line.strip() for line in chapters[0].raw_text.splitlines()}


def test_section_divider_not_parsed_as_chapter(tmp_path: Path) -> None:
    """A coarser section divider between numbered headings should not survive as a chapter."""

    docx_path = tmp_path / "section-divider.docx"
    _write_docx(
        docx_path,
        [
            ("A Sample Classic", "Title"),
            ("by Jane Doe", None),
            ("Introduction", "Heading 1"),
            ("Opening framing text.", "Body Text"),
            ("Chapter 1", "Heading 2"),
            ("First chapter body.", "Body Text"),
            ("Part Two", "Heading 1"),
            ("Sample Classic", "Normal"),
            ("Chapter 2", "Heading 2"),
            ("Second chapter body.", "Body Text"),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.title for chapter in chapters] == ["Introduction", "Chapter 1", "Chapter 2"]
    assert all(chapter.title != "Part Two" for chapter in chapters)


def test_primary_numbered_heading_level_prefers_most_populated_deeper_level(tmp_path: Path) -> None:
    """Deeper repeated chapter headings should beat shallower book wrappers."""

    docx_path = tmp_path / "primary-numbered-level.docx"
    _write_docx(
        docx_path,
        [
            ("A Tale in Layers", "Title"),
            ("by Jane Doe", None),
            ("Book the First", "Heading 1"),
            ("Chapter I", "Heading 2"),
            ("The first chapter body begins here.", None),
            ("Chapter II", "Heading 2"),
            ("The second chapter body continues the story.", None),
            ("Book the Second", "Heading 1"),
            ("Chapter III", "Heading 2"),
            ("The third chapter body closes the sequence.", None),
        ],
    )

    parser = DocxParser()
    heading_level = parser._detect_primary_numbered_heading_level(parser._collect_paragraphs(Document(docx_path)))
    _metadata, chapters = parser.parse(docx_path)

    assert heading_level == 2
    assert [chapter.title for chapter in chapters] == ["Chapter 1", "Chapter 2", "Chapter 3"]
    assert all(chapter.title not in {"Book the First", "Book the Second"} for chapter in chapters)


def test_introduction_does_not_absorb_numbered_title_headings(tmp_path: Path) -> None:
    """Real numbered title headings must break out of the introduction body."""

    docx_path = tmp_path / "intro-numbered-title-headings.docx"
    _write_docx(
        docx_path,
        [
            ("A Devotional Manual", "Title"),
            ("by Jane Doe", None),
            ("Introduction", "Heading 1"),
            ("The introduction frames the spiritual argument.", None),
            ("1. The Soul's Great Need", "Heading 2"),
            ("The first numbered chapter body begins here.", None),
            ("2. The Competitive Laws of Service", "Heading 2"),
            ("The second numbered chapter body follows.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.title for chapter in chapters] == [
        "Introduction",
        "The Soul's Great Need",
        "The Competitive Laws of Service",
    ]
    assert "The Soul's Great Need" not in chapters[0].raw_text
    assert chapters[1].number == 1
    assert chapters[2].number == 2


def test_part_headings_parse_as_primary_chapters(tmp_path: Path) -> None:
    """Part headings should become real chapters when they are the primary structure."""

    docx_path = tmp_path / "part-headings-primary.docx"
    _write_docx(
        docx_path,
        [
            ("Daisy Miller", "Title"),
            ("by Jane Doe", None),
            ("Part I", "Heading 1"),
            ("Winterbourne first meets Daisy in Vevey.", None),
            ("Part II", "Heading 1"),
            ("The Roman section concludes the novella.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.title for chapter in chapters] == ["Part I", "Part II"]
    assert [chapter.number for chapter in chapters] == [1, 2]


def test_reversed_ordinal_book_headings_parse_as_chapters(tmp_path: Path) -> None:
    """Ordinal-first book headings should parse as numbered chapters."""

    docx_path = tmp_path / "reversed-ordinal-book.docx"
    _write_docx(
        docx_path,
        [
            ("The Will to Power", "Title"),
            ("by Jane Doe", None),
            ("First Book - European Nihilism", "Heading 2"),
            ("The first section body begins here.", None),
            ("Second Book - Discipline and Breeding", "Heading 2"),
            ("The second section body follows here.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.title for chapter in chapters] == [
        "First Book - European Nihilism",
        "Second Book - Discipline and Breeding",
    ]
    assert [chapter.number for chapter in chapters] == [1, 2]


def test_maxim_headings_parse_as_chapters(tmp_path: Path) -> None:
    """Maxim headings should count as numbered chapters."""

    docx_path = tmp_path / "maxim-headings.docx"
    _write_docx(
        docx_path,
        [
            ("The Maxims of War", "Title"),
            ("by Jane Doe", None),
            ("Maxim I", "Heading 3"),
            ("Know the terrain before you commit your forces.", None),
            ("Maxim II", "Heading 3"),
            ("Preserve strength for the decisive encounter.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.title for chapter in chapters] == ["Maxim I", "Maxim II"]
    assert [chapter.number for chapter in chapters] == [1, 2]


def test_dominant_heading_level_fallback_parses_unnumbered_heading_series(tmp_path: Path) -> None:
    """When numbering is absent, the dominant heading level should drive chapter extraction."""

    docx_path = tmp_path / "dominant-heading-fallback.docx"
    paragraphs: list[tuple[str, str | None]] = [
        ("Letters from a Stoic", "Title"),
        ("by Jane Doe", None),
        ("Collection Overview", "Heading 2"),
        ("This editorial scaffolding should not become a narratable chapter.", None),
    ]
    for index in range(1, 16):
        paragraphs.extend(
            [
                (f"ON TOPIC {index}", "Heading 3"),
                (f"This is the body text for topic {index}, and it should remain attached to its own heading.", None),
            ]
        )
    _write_docx(docx_path, paragraphs)

    _metadata, chapters = DocxParser().parse(docx_path)

    assert len(chapters) == 15
    assert chapters[0].title == "ON TOPIC 1"
    assert chapters[-1].title == "ON TOPIC 15"
    assert all(chapter.title != "Collection Overview" for chapter in chapters)


def test_parse_as_a_man_thinketh_manuscript_preserves_mixed_hierarchy() -> None:
    """Book 5 should keep its mixed heading hierarchy without promoting subheads."""

    folder = _find_as_a_man_thinketh_folder()
    assert folder is not None, "As a Man Thinketh manuscript folder not found."

    metadata, chapters, manuscript_path = ManuscriptParserFactory.parse_manuscript(folder)

    assert metadata.title == "As a Man Thinketh & From Poverty"
    assert len(chapters) == 19
    assert manuscript_path.name == "005 As a Man Thinketh - From Poverty to Power-6x9.docx"
    assert [chapter.title for chapter in chapters[:8]] == [
        "Introduction",
        "THOUGHT AND CHARACTER",
        "EFFECT OF THOUGHT ON CIRCUMSTANCES",
        "EFFECT OF THOUGHT ON HEALTH AND THE BODY",
        "THOUGHT AND PURPOSE",
        "THE THOUGHT-FACTOR IN ACHIEVEMENT",
        "VISIONS AND IDEALS",
        "SERENITY",
    ]
    assert [chapter.title for chapter in chapters[8:11]] == [
        "The lesson of evil",
        "The world a reflex of mental states",
        "The way out of undesirable conditions",
    ]
    assert [chapter.title for chapter in chapters[-4:]] == [
        "The power of meditation",
        "The two masters, self and truth",
        "The acquirement of spiritual power",
        "The realization of perfect peace",
    ]
    assert "As a Man Thinketh" not in [chapter.title for chapter in chapters]
    assert "From Poverty to Power" not in [chapter.title for chapter in chapters]
    assert "Foreword" not in [chapter.title for chapter in chapters]
    assert "The realization of selfless love" not in [chapter.title for chapter in chapters]
    assert "Entering into the infinite" not in [chapter.title for chapter in chapters]
    assert "Saints, sages, and saviors: the law of service" not in [chapter.title for chapter in chapters]
    assert all(
        "As a Man Thinketh & From Poverty to Power" not in {line.strip() for line in chapter.raw_text.splitlines()}
        for chapter in chapters
    )
    assert all(
        not any(
            line.strip().startswith(("As a Man Thinketh:", "From Poverty to Power:"))
            for line in chapter.raw_text.splitlines()
        )
        for chapter in chapters
    )


def test_parse_nicomachean_ethics_keeps_all_ten_books() -> None:
    """Book 12 should not stop at the prose sentence 'The end is...' inside Book 3."""

    folder = _find_nicomachean_ethics_folder()
    assert folder is not None, "Nicomachean Ethics manuscript folder not found."

    metadata, chapters, manuscript_path = ManuscriptParserFactory.parse_manuscript(folder)

    assert metadata.title == "Nicomachean Ethics"
    assert manuscript_path.name == "011-Nicomachean-Ethics-Friedrich-Nietzsche-6x9-241.docx"
    assert [chapter.title for chapter in chapters] == [
        "Introduction",
        "Book 1",
        "Book 2",
        "Book 3",
        "Book 4",
        "Book 5",
        "Book 6",
        "Book 7",
        "Book 8",
        "Book 9",
        "Book 10",
    ]
    assert "The end is what we wish for" in chapters[3].raw_text
    assert chapters[4].raw_text.startswith("Let’s now talk about liberality.")


def test_parse_shortness_of_life_skips_author_and_message_to_reader() -> None:
    """Book 7 should skip author/title-page front matter and start at the real work."""

    folder = _find_shortness_of_life_folder()
    assert folder is not None, "On the Shortness of Life manuscript folder not found."

    metadata, chapters, manuscript_path = ManuscriptParserFactory.parse_manuscript(folder)

    assert metadata.title == "On the Shortness of Life"
    assert manuscript_path.name == "007 On the Shortness of Life-Seneca-5x8-021.docx"
    assert len(chapters) == 1
    assert chapters[0].title == "ON THE SHORTNESS OF LIFE"
    assert "Dear Reader" not in chapters[0].raw_text
    assert "MESSAGE TO THE READER" not in chapters[0].raw_text
    assert "Translated by Tim Zengerink" not in chapters[0].raw_text


def test_reader_note_front_matter_is_excluded(tmp_path: Path) -> None:
    """Reader-note sections before the first chapter should never be narrated."""

    docx_path = tmp_path / "reader-note-front-matter.docx"
    _write_docx(
        docx_path,
        [
            ("A Sample Treatise", "Title"),
            ("by Jane Doe", None),
            ("A Note to the Reader", "Heading 1"),
            ("This prefatory note should be skipped completely.", None),
            ("Chapter 1", "Heading 1"),
            ("This is the real narratable chapter body.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert len(chapters) == 1
    assert chapters[0].title == "Chapter 1"
    assert "A Note to the Reader" not in chapters[0].raw_text
    assert "prefatory note should be skipped" not in chapters[0].raw_text


def test_back_matter_outro_sections_are_excluded(tmp_path: Path) -> None:
    """Afterwords and author notes after the real chapters should terminate parsing."""

    docx_path = tmp_path / "back-matter-outro.docx"
    _write_docx(
        docx_path,
        [
            ("A Sample Treatise", "Title"),
            ("by Jane Doe", None),
            ("Chapter 1", "Heading 1"),
            ("This is the real narratable chapter body.", None),
            ("Author's Note", "Heading 1"),
            ("This end note should not be narrated.", None),
            ("Afterword", "Heading 1"),
            ("Neither should this afterword.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert len(chapters) == 1
    assert chapters[0].title == "Chapter 1"
    assert "Author's Note" not in chapters[0].raw_text
    assert "Afterword" not in chapters[0].raw_text
    assert "end note should not be narrated" not in chapters[0].raw_text


def test_reader_note_and_outro_patterns_are_globally_excluded() -> None:
    """Configured reader-note/outro markers should be skipped before start and terminate at the end."""

    parser = DocxParser()

    for text in (
        "Message to the Reader",
        "A Note to the Reader",
        "Note to the Reader",
        "Outro",
        "Closing Note",
        "Afterword",
        "A Word from the Author",
        "From the Author",
        "Author's Note",
    ):
        assert parser._is_skip_section(text)

    for text in (
        "Message to the Reader",
        "Note to the Reader",
        "Afterword",
        "Author's Note",
        "Closing Note",
        "From the Author",
    ):
        assert parser._is_back_matter_section(text)

    assert parser._is_skip_section("Foreword")
    assert parser._is_back_matter_section("The End")
    assert parser._is_back_matter_section("The End.") is True
    assert parser._is_back_matter_section("The end is what we wish for") is False


def test_standalone_chapter_marker_uses_following_heading_title(tmp_path: Path) -> None:
    """A plain Chapter N marker should yield the real title from the next heading-styled line."""

    docx_path = tmp_path / "standalone-marker-title.docx"
    _write_docx(
        docx_path,
        [
            ("A Warrior's Manual", "Title"),
            ("by Jane Doe", None),
            ("CHAPTER 1", "Normal"),
            ("THE GROUND BOOK", "Heading 2"),
            ("Strategy is the skill of the warrior.", "Body Text"),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert len(chapters) == 1
    assert chapters[0].number == 1
    assert chapters[0].title == "THE GROUND BOOK"
    assert chapters[0].raw_text == "Strategy is the skill of the warrior."


def test_parse_synthetic_docx_skip_rules_and_intro(tmp_path: Path) -> None:
    """Keep TOC entries and in-chapter subheads from becoming false chapter breaks."""

    document = Document()
    document.add_paragraph("A Sample Mystery", style="Title")
    document.add_paragraph("by Jane Doe")
    document.add_paragraph("Copyright 2026")
    document.add_paragraph("Table of Contents")
    document.add_paragraph("Chapter I. The Beginning")
    document.add_paragraph("Chapter II. The Middle")
    document.add_paragraph("Introduction", style="Heading 1")
    document.add_paragraph("Opening setup.")
    document.add_paragraph("Chapter I. The Beginning", style="Heading 1")
    document.add_paragraph("The story starts here.")
    document.add_paragraph("I.", style="Heading 2")
    document.add_paragraph("A subsection should remain inside chapter one.")
    document.add_paragraph("Thank You for Reading")

    docx_path = tmp_path / "sample.docx"
    document.save(docx_path)

    metadata, chapters = DocxParser().parse(docx_path)

    assert metadata.title == "A Sample Mystery"
    assert metadata.author == "Jane Doe"
    assert [chapter.type for chapter in chapters] == ["introduction", "chapter"]
    assert chapters[0].number == 0
    assert chapters[1].number == 1
    assert chapters[1].title == "The Beginning"
    assert "A subsection should remain inside chapter one." in chapters[1].raw_text
    assert "Thank You for Reading" not in chapters[1].raw_text


def test_introduction_with_subheadings(tmp_path: Path) -> None:
    """Sub-headings inside an introduction should stay inside the introduction body."""

    docx_path = tmp_path / "intro-subheadings.docx"
    _write_docx(
        docx_path,
        [
            ("The Art of Strategy", "Title"),
            ("by Jane Doe", None),
            ("Introduction", "Heading 1"),
            ("I. Biography", "Heading 3"),
            ("The introduction opens with a biographical overview of the author and the world that shaped the text.", None),
            ("II. Historical Context", "Heading 3"),
            ("A second introductory section explains the political pressures, strategic traditions, and conflicts surrounding the work.", None),
            ("Chapter 1 - The Beginning", "Heading 1"),
            ("The first chapter body begins here.", None),
            ("Chapter 2 - The Middle", "Heading 1"),
            ("The second chapter body follows.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.type for chapter in chapters] == ["introduction", "chapter", "chapter"]
    assert chapters[0].title == "Introduction"
    assert "I. Biography" in chapters[0].raw_text
    assert "II. Historical Context" in chapters[0].raw_text
    assert "biographical overview" in chapters[0].raw_text
    assert "political pressures" in chapters[0].raw_text
    assert chapters[1].number == 1
    assert chapters[1].title == "The Beginning"
    assert chapters[2].number == 2
    assert chapters[2].title == "The Middle"


def test_introduction_without_subheadings(tmp_path: Path) -> None:
    """A simple introduction should still parse exactly as before."""

    docx_path = tmp_path / "intro-regression.docx"
    _write_docx(
        docx_path,
        [
            ("A Sample Treatise", "Title"),
            ("by Jane Doe", None),
            ("Introduction", "Heading 1"),
            ("This introduction provides context, stakes, and framing before the main numbered chapters begin.", None),
            ("Chapter 1 - The Beginning", "Heading 1"),
            ("The first chapter body begins here.", None),
            ("Chapter 2 - The Middle", "Heading 1"),
            ("The second chapter body follows.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.type for chapter in chapters] == ["introduction", "chapter", "chapter"]
    assert chapters[0].title == "Introduction"
    assert "context, stakes, and framing" in chapters[0].raw_text
    assert chapters[1].title == "The Beginning"
    assert chapters[2].title == "The Middle"


def test_numbered_book_hierarchy_blocks_deeper_false_chapter_splits(tmp_path: Path) -> None:
    """A dominant top-level book sequence should suppress deeper numbered subheadings."""

    docx_path = tmp_path / "book-hierarchy.docx"
    _write_docx(
        docx_path,
        [
            ("Meditations", "Title"),
            ("Marcus Aurelius", None),
            ("Introduction", "Heading 1"),
            ("Opening framing for the text.", None),
            ("1. Daily Reflection and Journaling", "Heading 4"),
            ("This should remain inside the introduction body.", None),
            ("Book One", "Heading 1"),
            ("From my grandfather Verus, I learned kindness.", None),
            ("Book Two", "Heading 1"),
            ("Begin the day by reminding yourself what people can be like.", None),
            ("Book Three", "Heading 1"),
            ("Do not waste what remains of your life in idle thoughts.", None),
            ("Chapter 16", "Heading 2-Intro"),
            ("This stray subheading belongs inside Book Three.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.title for chapter in chapters] == [
        "Introduction",
        "Book One",
        "Book Two",
        "Book Three",
    ]
    assert "Daily Reflection and Journaling" in chapters[0].raw_text
    assert "Chapter 16" in chapters[3].raw_text


def test_intro_followed_by_heading_one_story_title_starts_real_chapter(tmp_path: Path) -> None:
    """A Heading 1 story title after an introduction should not be absorbed into the intro."""

    docx_path = tmp_path / "short-story-title-heading.docx"
    _write_docx(
        docx_path,
        [
            ("The Moon-Bog", "Title"),
            ("A Lovecraftian Tale of Ancient Curses", "Subtitle-1-Book"),
            ("H.P. Lovecraft", "Subtitle-3-Author"),
            ("Table of Contents", "TOC Heading"),
            ("Introduction 1", "TOC 1"),
            ("The Moon-Bog 9", "TOC 1"),
            ("Introduction", "Heading 1"),
            ("This introduction explains the story context and historical setting.", None),
            ("The Moon-Bog", "Heading 1"),
            ("No one knows exactly where Denys Barry went.", None),
        ],
    )

    metadata, chapters = DocxParser().parse(docx_path)

    assert metadata.author == "H.P. Lovecraft"
    assert metadata.subtitle == "A Lovecraftian Tale of Ancient Curses"
    assert [chapter.type for chapter in chapters] == ["introduction", "chapter"]
    assert chapters[0].title == "Introduction"
    assert chapters[1].title == "The Moon-Bog"
    assert "No one knows exactly where Denys Barry went." in chapters[1].raw_text


def test_heading_one_story_title_after_prologue_becomes_chapter(tmp_path: Path) -> None:
    """Unnumbered Heading 1 story sections should follow a prologue as real chapters."""

    docx_path = tmp_path / "prologue-then-story-title.docx"
    _write_docx(
        docx_path,
        [
            ("The White People", "Title"),
            ("Arthur Machen", "Subtitle-3-Author"),
            ("Prologue", "Heading 1"),
            ("Magic and holiness are the only real things.", None),
            ("The Green Book", "Heading 1"),
            ("I saw the white people in the woods.", None),
        ],
    )

    _metadata, chapters = DocxParser().parse(docx_path)

    assert [chapter.type for chapter in chapters] == ["introduction", "chapter"]
    assert chapters[0].title == "Prologue"
    assert chapters[1].number == 1
    assert chapters[1].title == "The Green Book"


def test_explicit_chapter_heading_detection() -> None:
    """Only real 'Chapter N' headings should count as explicit chapter boundaries."""

    parser = DocxParser()

    assert parser._is_explicit_chapter_heading("Chapter 1 - Laying Plans") is True
    assert parser._is_explicit_chapter_heading("Chapter IV") is True
    assert parser._is_explicit_chapter_heading("I. Brief Biography") is False
    assert parser._is_explicit_chapter_heading("Historical Context") is False
    assert parser._is_explicit_chapter_heading("1: Some Title") is False


def test_build_chapter_returns_none_for_empty_body() -> None:
    """Empty section bodies should be skipped instead of raising exceptions."""

    parser = DocxParser()

    assert parser._build_chapter({"number": 1, "title": "Chapter 1", "type": "chapter"}, ["", "   "]) is None


def test_parse_skips_empty_chapter_sections_instead_of_crashing(tmp_path: Path) -> None:
    """A heading without body text should be omitted while later chapters still parse."""

    docx_path = tmp_path / "empty-chapter.docx"
    _write_docx(
        docx_path,
        [
            ("A Stoic Handbook", "Title"),
            ("by Jane Doe", None),
            ("Chapter I. Empty Start", "Heading 1"),
            ("Chapter II. Real Chapter", "Heading 1"),
            ("This chapter has actual body text.", None),
        ],
    )

    metadata, chapters = DocxParser().parse(docx_path)

    assert metadata.author == "Jane Doe"
    assert [chapter.number for chapter in chapters] == [2]
    assert chapters[0].title == "Real Chapter"
    assert chapters[0].raw_text == "This chapter has actual body text."


def test_parse_skips_long_toc_entries_and_finds_real_chapters(tmp_path: Path) -> None:
    """Long TOC entries with TOC styles should not become empty chapters."""

    long_chapter_one = (
        "Chapter I: In Which the North Polar Practical Association Decides That It Is "
        "Absolutely Necessary to Reach the Pole"
    )
    long_chapter_two = (
        "Chapter II: In Which the Delegation Is Sent Forth To Investigate the Frozen "
        "Regions and Report Back Quickly"
    )
    docx_path = tmp_path / "long-toc.docx"
    _write_docx(
        docx_path,
        [
            ("An Arctic Adventure", "Title"),
            ("by Jane Doe", None),
            ("Table of Contents", None),
            (long_chapter_one, "TOC 1"),
            (long_chapter_two, "TOC 1"),
            (long_chapter_one, "Heading 1"),
            ("The expedition begins in earnest.", None),
            (long_chapter_two, "Heading 1"),
            ("The report confirms the danger ahead.", None),
        ],
    )

    parser = DocxParser()
    _metadata, chapters = parser.parse(docx_path)

    assert [chapter.number for chapter in chapters] == [1, 2]
    assert chapters[0].title == (
        "In Which the North Polar Practical Association Decides That It Is Absolutely Necessary to Reach the Pole"
    )
    assert chapters[1].title == (
        "In Which the Delegation Is Sent Forth To Investigate the Frozen Regions and Report Back Quickly"
    )
    assert parser.last_toc_entries == [long_chapter_one, long_chapter_two]


def test_parse_keeps_collecting_toc_for_toc_styled_note_paragraphs(tmp_path: Path) -> None:
    """TOC-styled note-like lines should not force an early exit from TOC mode."""

    docx_path = tmp_path / "toc-note.docx"
    _write_docx(
        docx_path,
        [
            ("Collected Speeches", "Title"),
            ("by Jane Doe", None),
            ("Table of Contents", None),
            ("Volume I", "TOC 1"),
            ("Chapter I. Opening Address", "TOC 1"),
            ("Chapter I. Opening Address", "Heading 1"),
            ("Actual chapter body text.", None),
        ],
    )

    parser = DocxParser()
    _metadata, chapters = parser.parse(docx_path)

    assert [chapter.number for chapter in chapters] == [1]
    assert parser.last_toc_entries == ["Volume I", "Chapter I. Opening Address"]


def test_toc_entries_with_page_numbers_are_detected() -> None:
    """TOC entries with trailing page numbers should be recognized after normalization."""

    parser = DocxParser()

    assert parser._looks_like_toc_entry("Chapter II. The Middle\t42", None) is True
    assert parser._looks_like_toc_entry("Chapter II. The Middle 42", None) is True


def test_modern_translation_line_is_not_treated_as_author(tmp_path: Path) -> None:
    """Edition descriptors should not beat a real author-like line."""

    docx_path = tmp_path / "author-false-positive.docx"
    _write_docx(
        docx_path,
        [
            ("Ancient Wisdom", "Title"),
            ("A Modern Translation", None),
            ("Jane Doe", None),
            ("Chapter I. Opening", "Heading 1"),
            ("Body text.", None),
        ],
    )

    metadata, _chapters = DocxParser().parse(docx_path)

    assert metadata.author == "Jane Doe"
    assert metadata.subtitle is None


def test_parse_falls_back_to_unknown_author_without_crashing(tmp_path: Path) -> None:
    """Missing author metadata should return a fallback instead of raising."""

    docx_path = tmp_path / "unknown-author.docx"
    _write_docx(
        docx_path,
        [
            ("Meditations on Resilience", "Title"),
            ("Chapter I. Opening", "Heading 1"),
            ("This manuscript omits the author line.", None),
        ],
    )

    metadata, chapters = DocxParser().parse(docx_path)

    assert metadata.author == "Unknown Author"
    assert len(chapters) == 1
    assert chapters[0].title == "Opening"


def test_fallback_single_chapter(tmp_path: Path) -> None:
    """Books without headings should fall back to one Full Text chapter."""

    body_opening = (
        "This manuscript begins immediately with sustained body prose that exceeds thirty words, "
        "so the parser should treat it as the start of narratable content instead of discarding "
        "the whole document for lacking explicit chapter headings."
    )
    body_followup = (
        "A second paragraph should stay attached to the same fallback chapter, preserving spacing "
        "and ensuring the parser emits one continuous narratable section for the full document."
    )
    docx_path = tmp_path / "full-text-fallback.docx"
    _write_docx(
        docx_path,
        [
            ("Dialogues on Strategy", "Title"),
            ("by Jane Doe", None),
            ("Copyright 2026", None),
            (body_opening, None),
            (body_followup, None),
            ("Thank You for Reading", None),
        ],
    )

    metadata, chapters = DocxParser().parse(docx_path)

    assert metadata.author == "Jane Doe"
    assert len(chapters) == 1
    assert chapters[0].number == 1
    assert chapters[0].title == "Full Text"
    assert chapters[0].type == "chapter"
    assert chapters[0].raw_text == f"{body_opening}\n\n{body_followup}"


def test_fallback_empty_document(tmp_path: Path) -> None:
    """Front-matter-only documents should still fail with diagnostics."""

    docx_path = tmp_path / "front-matter-only.docx"
    _write_docx(
        docx_path,
        [
            ("A Book Without Chapters", "Title"),
            ("by Jane Doe", None),
            ("Copyright 2026", None),
            ("Table of Contents", None),
            ("Thank You for Reading", None),
        ],
    )

    with pytest.raises(ValueError, match=r"No narratable chapters detected in front-matter-only\.docx"):
        DocxParser().parse(docx_path)


def test_extract_author_from_folder_name() -> None:
    """Known-author folder patterns should resolve to canonical author names."""

    parser = DocxParser()

    assert parser._extract_author_from_folder("001-The-Art-of-War-Sun-Tzu-6x9-142") == "Sun Tzu"
    assert (
        parser._extract_author_from_folder("010-Beyond-Good-and-Evil-Friedrich-Nietzsche-6x9-262")
        == "Friedrich Nietzsche"
    )
    assert parser._extract_author_from_folder("020-The-Kybalion-6x9-113") is None


def test_factory_uses_folder_hint_for_unknown_docx_authors(tmp_path: Path) -> None:
    """Factory parsing should upgrade Unknown Author when the folder embeds a known author."""

    manuscript_folder = tmp_path / "001-The-Art-of-War-Sun-Tzu-6x9-142"
    manuscript_folder.mkdir()
    docx_path = manuscript_folder / "manuscript.docx"
    _write_docx(
        docx_path,
        [
            ("The Art of War", "Title"),
            ("Chapter I. Laying Plans", "Heading 1"),
            ("Appear weak when you are strong.", None),
        ],
    )

    metadata, chapters, manuscript_path = ManuscriptParserFactory.parse_manuscript(manuscript_folder)

    assert metadata.author == "Sun Tzu"
    assert len(chapters) == 1
    assert manuscript_path == docx_path


def test_parse_with_folder_hint_uses_title_lookup_for_unknown_author(tmp_path: Path) -> None:
    """Known titles should resolve Unknown Author when the folder name has no author hint."""

    docx_path = tmp_path / "known-title.docx"
    _write_docx(
        docx_path,
        [
            ("Metaphysics", "Title"),
            ("Chapter I. First Principles", "Heading 1"),
            ("All men by nature desire to know.", None),
        ],
    )

    metadata, chapters = DocxParser().parse_with_folder_hint(docx_path, folder_name="001-Metaphysics-6x9-250")

    assert metadata.author == "Aristotle"
    assert len(chapters) == 1


def test_known_titles_expansion() -> None:
    """Expanded title mappings should cover the new canonical works."""

    parser = DocxParser()

    assert parser.KNOWN_TITLES["jane eyre"] == "Charlotte Brontë"
    assert parser.KNOWN_TITLES["don quixote"] == "Miguel de Cervantes"
    assert parser.KNOWN_TITLES["the communist manifesto"] == "Karl Marx"
    assert parser.KNOWN_TITLES["the epic of gilgamesh"] == "Anonymous"


def test_known_authors_expansion() -> None:
    """Expanded folder-author mappings should resolve the new canonical authors."""

    parser = DocxParser()

    assert parser.KNOWN_AUTHORS["xenophon"] == "Xenophon"
    assert parser.KNOWN_AUTHORS["charlotte brontë"] == "Charlotte Brontë"


def test_title_lookup_with_ampersand(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Title lookup should normalize ampersands to 'and' before failing."""

    monkeypatch.setattr(
        DocxParser,
        "KNOWN_TITLES",
        {"the complete strategy and war collection": "Various Authors"},
    )

    docx_path = tmp_path / "ampersand-title.docx"
    _write_docx(
        docx_path,
        [
            ("The Complete Strategy & War Collection", "Title"),
            ("Chapter I. Opening", "Heading 1"),
            ("Collected strategic writings begin here.", None),
        ],
    )

    metadata, _chapters = DocxParser().parse_with_folder_hint(
        docx_path,
        folder_name="001-Complete-Strategy-Collection-6x9-250",
    )

    assert metadata.author == "Various Authors"


def test_folder_author_hint_takes_priority_over_title_lookup(tmp_path: Path) -> None:
    """Folder-name author extraction should win over title-based fallback."""

    docx_path = tmp_path / "known-title-priority.docx"
    _write_docx(
        docx_path,
        [
            ("Metaphysics", "Title"),
            ("Chapter I. First Principles", "Heading 1"),
            ("All men by nature desire to know.", None),
        ],
    )

    metadata, _chapters = DocxParser().parse_with_folder_hint(
        docx_path,
        folder_name="001-Metaphysics-Sun-Tzu-6x9-250",
    )

    assert metadata.author == "Sun Tzu"


def test_non_author_phrases_are_rejected_as_author_names() -> None:
    """Known subtitle and edition phrases should never pass the author heuristic."""

    parser = DocxParser()

    for phrase in parser.NON_AUTHOR_PHRASES:
        assert parser._looks_like_author_name(phrase.title()) is False


def test_no_narratable_chapters_error_includes_diagnostics(tmp_path: Path) -> None:
    """The parser should explain why no chapters were found."""

    docx_path = tmp_path / "no-chapters.docx"
    _write_docx(
        docx_path,
        [
            ("A Book Without Headings", "Title"),
            ("by Jane Doe", None),
            ("This file has body text.", None),
            ("But no narratable chapter markers.", None),
        ],
    )

    with pytest.raises(ValueError, match=r"No narratable chapters detected in no-chapters\.docx") as exc_info:
        DocxParser().parse(docx_path)

    message = str(exc_info.value)
    assert "The document has 4 paragraphs (4 non-empty)." in message
    assert "The parser requires chapter headings with 'Chapter N' format or Heading-styled paragraphs." in message


def test_skip_rules_ignore_case_and_punctuation_variants() -> None:
    """Skip rules should match front and back matter despite punctuation changes."""

    assert should_skip_heading("PREFACE - Message to the Reader")
    assert should_skip_heading("Thank You For Reading!!!")


def test_text_cleaning() -> None:
    """Normalize page markers, abbreviations, and punctuation for TTS."""

    cleaner = TextCleaner()
    dirty_text = (
        "Page 42\n"
        "Dr. Watson and Mr. Holmes discussed the case—a difficult one…\n"
        "They met on Baker St. near St. Paul's.\n"
        "Page 43"
    )

    cleaned_text = cleaner.clean(dirty_text)

    assert "Page 42" not in cleaned_text
    assert "Page 43" not in cleaned_text
    assert "Doctor Watson" in cleaned_text
    assert "Mister Holmes" in cleaned_text
    assert "Baker Street" in cleaned_text
    assert "Saint Paul's" in cleaned_text
    assert "..." in cleaned_text
    assert "—" in cleaned_text


def test_credits_generation() -> None:
    """Generate opening and closing credits with the required metadata."""

    opening = CreditsGenerator.generate_opening_credits(
        title="The Sherlock Holmes Mysteries",
        subtitle="A Complete Collection",
        author="Arthur Conan Doyle",
    )
    closing = CreditsGenerator.generate_closing_credits(
        title="The Sherlock Holmes Mysteries",
        subtitle="A Complete Collection",
        author="Arthur Conan Doyle",
    )

    assert "This is The Sherlock Holmes Mysteries." in opening
    assert "A Complete Collection." in opening
    assert "Written by Arthur Conan Doyle." in opening
    assert "Narrated by James Mitchell." in opening

    assert "This was The Sherlock Holmes Mysteries." in closing
    assert "A Complete Collection." in closing
    assert "Written by Arthur Conan Doyle." in closing
    assert "Narrated by James Mitchell." in closing
