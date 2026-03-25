"""Helpers for generating parser fixture manuscripts in multiple formats."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from ebooklib import epub


@dataclass(slots=True)
class SectionSpec:
    """Structured section definition shared by parser tests."""

    heading: str
    paragraphs: list[str]
    heading_tag: str = "h1"


def default_book_sections() -> list[SectionSpec]:
    """Return a canonical set of narratable sections for cross-format tests."""

    return [
        SectionSpec(
            heading="Introduction",
            paragraphs=[
                "Opening setup for the mystery.",
                "The narrator meets the case in a dim library.",
            ],
        ),
        SectionSpec(
            heading="Chapter I. The Beginning",
            paragraphs=[
                "The story starts here with a suspicious visitor.",
                "A clue appears in the hallway before dawn.",
            ],
        ),
        SectionSpec(
            heading="Chapter II. The Twist",
            paragraphs=[
                "The case becomes stranger with each new witness.",
                "Nothing is what it first seemed to be.",
            ],
        ),
    ]


def create_sample_docx(
    path: Path,
    *,
    title: str,
    subtitle: str | None,
    author: str,
    sections: list[SectionSpec],
) -> None:
    """Create a DOCX manuscript with front matter, TOC, and narratable sections."""

    document = Document()
    document.add_paragraph(title, style="Title")
    if subtitle:
        document.add_paragraph(subtitle)
    document.add_paragraph(f"by {author}")
    document.add_paragraph("Copyright 2026")
    document.add_paragraph("Table of Contents")
    for section in sections:
        document.add_paragraph(section.heading)

    for section in sections:
        document.add_paragraph(section.heading, style="Heading 1")
        for paragraph in section.paragraphs:
            document.add_paragraph(paragraph)

    document.add_paragraph("Thank You for Reading")
    document.save(path)


def create_sample_epub(
    path: Path,
    *,
    title: str,
    subtitle: str | None,
    author: str,
    sections: list[SectionSpec],
) -> None:
    """Create an EPUB manuscript with front matter and one XHTML document per section."""

    book = epub.EpubBook()
    book.set_identifier("parser-test-book")
    book.set_title(title)
    book.set_language("en")
    book.add_author(author)
    if subtitle:
        book.add_metadata("DC", "description", subtitle)

    items: list[epub.EpubHtml] = []

    title_page = epub.EpubHtml(title="Title Page", file_name="title.xhtml", lang="en")
    title_page.content = f"""
    <html><body>
      <h1>Title Page</h1>
      <p>{title}</p>
      <p>{subtitle or ''}</p>
      <p>by {author}</p>
    </body></html>
    """
    book.add_item(title_page)
    items.append(title_page)

    toc_page = epub.EpubHtml(title="Table of Contents", file_name="toc.xhtml", lang="en")
    toc_entries = "".join(f"<p>{section.heading}</p>" for section in sections)
    toc_page.content = f"<html><body><h1>Table of Contents</h1>{toc_entries}</body></html>"
    book.add_item(toc_page)
    items.append(toc_page)

    for index, section in enumerate(sections, start=1):
        chapter = epub.EpubHtml(title=section.heading, file_name=f"chapter-{index}.xhtml", lang="en")
        paragraphs = "".join(f"<p>{paragraph}</p>" for paragraph in section.paragraphs)
        chapter.content = (
            "<html><body>"
            f"<{section.heading_tag}>{section.heading}</{section.heading_tag}>"
            f"{paragraphs}"
            "</body></html>"
        )
        book.add_item(chapter)
        items.append(chapter)

    back_matter = epub.EpubHtml(title="Back Matter", file_name="thanks.xhtml", lang="en")
    back_matter.content = "<html><body><h1>Thank You for Reading</h1><p>Skipped back matter.</p></body></html>"
    book.add_item(back_matter)
    items.append(back_matter)

    book.toc = tuple(items)
    book.spine = ["nav", *items]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    epub.write_epub(str(path), book)


def create_sample_pdf(
    path: Path,
    *,
    title: str,
    subtitle: str | None,
    author: str,
    sections: list[SectionSpec],
    include_metadata: bool = True,
    extra_pages: int = 0,
) -> None:
    """Create a text-based PDF manuscript suitable for pdfplumber extraction."""

    pages: list[list[str]] = [
        [title, *( [subtitle] if subtitle else [] ), f"by {author}"],
        ["Copyright 2026", "Table of Contents", *[section.heading for section in sections]],
    ]

    for section in sections:
        pages.append([section.heading, *section.paragraphs])

    for page_number in range(extra_pages):
        pages.append([f"Continuation page {page_number + 1} for large document coverage."])

    pages.append(["Thank You for Reading", "Skipped back matter."])

    metadata = {"Title": title, "Author": author, "Subject": subtitle or ""} if include_metadata else None
    write_text_pdf(path, pages, metadata=metadata)


def create_consistency_fixture(folder: Path) -> dict[str, Path]:
    """Generate the same sample manuscript in DOCX, EPUB, and PDF formats."""

    sections = default_book_sections()
    title = "The Test Chronicle"
    subtitle = "A Detective Story"
    author = "Jane Doe"

    docx_path = folder / "book.docx"
    epub_path = folder / "book.epub"
    pdf_path = folder / "book.pdf"

    create_sample_docx(docx_path, title=title, subtitle=subtitle, author=author, sections=sections)
    create_sample_epub(epub_path, title=title, subtitle=subtitle, author=author, sections=sections)
    create_sample_pdf(pdf_path, title=title, subtitle=subtitle, author=author, sections=sections)

    return {
        "docx": docx_path,
        "epub": epub_path,
        "pdf": pdf_path,
    }


def write_text_pdf(path: Path, pages: list[list[str]], metadata: dict[str, str] | None = None) -> None:
    """Write a minimal text PDF with searchable Helvetica text."""

    objects: list[bytes] = []

    def add_object(body: str) -> int:
        objects.append(body.encode("latin-1"))
        return len(objects)

    font_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    page_ids: list[int] = []

    for lines in pages:
        stream = _build_page_stream(lines)
        content_id = add_object(
            f"<< /Length {len(stream)} >>\nstream\n{stream.decode('latin-1')}\nendstream"
        )
        page_id = add_object(
            "<< /Type /Page /Parent 0 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>"
        )
        page_ids.append(page_id)

    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    pages_id = add_object(f"<< /Type /Pages /Count {len(page_ids)} /Kids [{kids}] >>")
    for page_id in page_ids:
        page_body = objects[page_id - 1].decode("latin-1").replace("/Parent 0 0 R", f"/Parent {pages_id} 0 R")
        objects[page_id - 1] = page_body.encode("latin-1")

    catalog_id = add_object(f"<< /Type /Catalog /Pages {pages_id} 0 R >>")
    info_id = None
    if metadata:
        escaped_pairs = " ".join(f"/{key} ({_escape_pdf_text(value)})" for key, value in metadata.items())
        info_id = add_object(f"<< {escaped_pairs} >>")

    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    offsets: list[int] = []
    body = bytearray(header)
    for object_id, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body.extend(f"{object_id} 0 obj\n".encode("latin-1"))
        body.extend(obj)
        body.extend(b"\nendobj\n")

    xref_start = len(body)
    body.extend(f"xref\n0 {len(objects) + 1}\n".encode("latin-1"))
    body.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        body.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))

    trailer = f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_id} 0 R"
    if info_id is not None:
        trailer += f" /Info {info_id} 0 R"
    trailer += f" >>\nstartxref\n{xref_start}\n%%EOF\n"
    body.extend(trailer.encode("latin-1"))
    path.write_bytes(bytes(body))


def _build_page_stream(lines: list[str]) -> bytes:
    """Build a PDF text stream for a list of lines."""

    operations = ["BT", "/F1 14 Tf", "72 720 Td"]
    first_line = True
    for line in lines:
        if not first_line:
            operations.append("0 -20 Td")
        first_line = False
        operations.append(f"({_escape_pdf_text(line)}) Tj")
    operations.append("ET")
    return "\n".join(operations).encode("latin-1")


def _escape_pdf_text(value: str) -> str:
    """Escape PDF string literals."""

    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
